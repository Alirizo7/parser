from __future__ import annotations

import tempfile
import unittest
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH

from attestation.services.extract import parse_perechen_positions_6_5
from attestation.services.render import render_6_5
from attestation.services.selfcheck import compare_6_5


def _merge_row(row, text: str) -> None:
    merged = row.cells[0]
    for cell in row.cells[1:]:
        merged = merged.merge(cell)
    merged.text = text


def _merge_pair(row, left: int, right: int, text: str) -> None:
    row.cells[left].merge(row.cells[right]).text = text


def _company() -> dict:
    return {
        "name": "SINOV KORXONASI",
        "parent": "BOSH TASHKILOT",
        "address": "Toshkent shahri",
        "product": "Sinov mahsuloti",
        "stir": "123456789",
        "ifut": "12345",
        "mxbt": "67890",
    }


def _workplace(no: str, **extra) -> dict:
    record = {
        "workplace_no": no,
        "position": "Kartadagi lavozim",
        "subdivision": "Kartadagi bo‘lim",
        "factors": {},
        "benefits": {},
        "ppe_provided": "yo‘q",
        "privileged_pension": "yo‘q",
    }
    record.update(extra)
    return record


class Perechen65Tests(unittest.TestCase):
    def test_reads_unique_physical_cells_after_old_doc_horizontal_merges(self):
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "perechen.docx"
            doc = Document()
            table = doc.add_table(rows=5, cols=6)
            table.rows[0].cells[0].text = "иш ўрнининг т/р"
            table.rows[0].cells[1].text = "касб, лавозим"
            table.rows[0].cells[2].text = "коди"
            _merge_row(table.rows[1], "Маъмурият")

            # После конвертации старого .doc row.cells выглядит как
            # [wp, wp, position, position, code, code].
            _merge_pair(table.rows[2], 0, 1, "000001")
            _merge_pair(table.rows[2], 2, 3, "Бош врач")
            _merge_pair(table.rows[2], 4, 5, "1120")

            _merge_row(table.rows[3], "Лаборатория")
            table.rows[4].cells[0].text = "000002a"
            table.rows[4].cells[1].text = "Лаборант"
            table.rows[4].cells[2].text = "3212"
            doc.save(path)

            positions, warnings = parse_perechen_positions_6_5(path)

        self.assertEqual(warnings, [])
        self.assertEqual([p["workplace_no"] for p in positions], ["000001", "000002а"])
        self.assertEqual(positions[0]["position"], "Бош врач")
        self.assertEqual(positions[0]["job_code"], "1120")
        self.assertEqual(positions[0]["subdivision"], "Маъмурият")
        self.assertEqual(positions[1]["subdivision"], "Лаборатория")
        self.assertEqual([p["order"] for p in positions], [0, 1])


class Render65Tests(unittest.TestCase):
    def test_render_uses_only_6_5_perechen_fields_order_and_reference_format(self):
        workplaces = [
            _workplace(
                "000002",
                position_from_perechen="Бошқа ҳужжатдаги лавозим",
                job_code="9999",
                position_from_perechen_6_5="Лаборант",
                job_code_6_5="3212",
                subdivision_6_5="Лаборатория",
                perechen_order_6_5=1,
            ),
            _workplace(
                "000001",
                position_from_perechen="Бошқа ҳужжатдаги лавозим",
                job_code="8888",
                position_from_perechen_6_5="Бош врач",
                job_code_6_5="1120",
                subdivision_6_5="Маъмурият",
                perechen_order_6_5=0,
            ),
        ]

        with tempfile.TemporaryDirectory() as tmp:
            output = render_6_5(_company(), workplaces, Path(tmp) / "6_5.docx")
            doc = Document(output)

        summary = doc.tables[1]
        self.assertEqual(summary.rows[3].cells[0].text, "Маъмурият")
        self.assertEqual(summary.rows[4].cells[0].text, "000001")
        self.assertEqual(summary.rows[4].cells[1].text, "Бош врач")
        self.assertEqual(summary.rows[4].cells[2].text, "1120")
        self.assertEqual(summary.rows[5].cells[0].text, "Лаборатория")
        self.assertEqual(summary.rows[6].cells[0].text, "000002")
        self.assertIs(summary.rows[4].cells[22]._tc, summary.rows[4].cells[23]._tc)

        data_row = summary.rows[4]
        self.assertEqual(data_row.cells[0].vertical_alignment, WD_CELL_VERTICAL_ALIGNMENT.CENTER)
        self.assertEqual(data_row.cells[1].vertical_alignment, WD_CELL_VERTICAL_ALIGNMENT.CENTER)
        self.assertEqual(data_row.cells[0].paragraphs[0].alignment, WD_ALIGN_PARAGRAPH.CENTER)
        self.assertEqual(data_row.cells[1].paragraphs[0].alignment, WD_ALIGN_PARAGRAPH.LEFT)
        self.assertTrue(all(run.bold is False for cell in data_row.cells
                            for paragraph in cell.paragraphs for run in paragraph.runs if run.text))
        self.assertTrue(data_row._tr.xpath("./w:trPr/w:cantSplit"))
        self.assertTrue(summary.rows[3]._tr.xpath("./w:trPr/w:cantSplit"))
        self.assertTrue(summary.rows[3].cells[0].paragraphs[0]._p.xpath("./w:pPr/w:keepNext"))
        body_children = list(doc._element.body.iterchildren())
        summary_index = body_children.index(summary._tbl)
        self.assertEqual(
            [child.tag.rsplit("}", 1)[-1] for child in body_children[summary_index + 1:]],
            ["sectPr"],
        )
        self.assertFalse(any(rel.reltype.endswith("/image") for rel in doc.part.rels.values()))

    def test_selfcheck_treats_latin_and_cyrillic_a_suffix_as_same_workplace(self):
        common = dict(
            position_from_perechen_6_5="Оператор",
            job_code_6_5="8111",
            subdivision_6_5="Ишчилар",
            perechen_order_6_5=0,
        )
        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            generated = render_6_5(
                _company(), [_workplace("000002а", **common)], tmp_path / "generated.docx"
            )
            reference = render_6_5(
                _company(), [_workplace("000002a", **common)], tmp_path / "reference.docx"
            )
            result = compare_6_5(generated, reference)

        self.assertTrue(result.ok, result.mismatches)
        self.assertEqual(result.mismatches, [])


if __name__ == "__main__":
    unittest.main()
