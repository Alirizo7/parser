from __future__ import annotations

import tempfile
import unittest
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from attestation.services.extract import parse_perechen, parse_perechen_positions_6_4
from attestation.services.pipeline import _repair_duplicate_workplace_numbers_from_perechen
from attestation.services.render import render_6_4


def _merge_row(row, text: str) -> None:
    merged = row.cells[0]
    for cell in row.cells[1:]:
        merged = merged.merge(cell)
    merged.text = text


class Perechen64Tests(unittest.TestCase):
    def test_hierarchy_order_and_repeated_physical_groups_are_preserved(self):
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "perechen.docx"
            doc = Document()
            table = doc.add_table(rows=8, cols=5)
            table.rows[0].cells[0].text = "иш ўрнининг т/р"
            table.rows[0].cells[1].text = "касб, лавозим"
            table.rows[0].cells[2].text = "коди"
            table.rows[0].cells[3].text = "ишловчилар сони"
            table.rows[0].cells[4].text = "шундан аёллар"
            _merge_row(table.rows[1], "PARENT")
            _merge_row(table.rows[2], "CHILD")
            table.rows[3].cells[0].text = "000001"
            table.rows[3].cells[1].text = "Operator"
            table.rows[3].cells[2].text = "111"
            table.rows[3].cells[3].text = "2"
            table.rows[3].cells[4].text = "1"
            table.rows[4].cells[0].text = "000002"
            table.rows[4].cells[1].text = "Muhandis"
            table.rows[4].cells[2].text = "222"
            table.rows[4].cells[3].text = "1"
            table.rows[4].cells[4].text = "-"
            _merge_row(table.rows[5], "PARENT")
            table.rows[6].cells[0].text = "000003"
            table.rows[6].cells[1].text = "Laborant"
            table.rows[6].cells[2].text = "333"
            table.rows[6].cells[3].text = "3"
            table.rows[6].cells[4].text = "2"
            # Номер приклеен к тексту из-за дефекта старого .doc.
            table.rows[7].cells[0].text = "PARENT000004"
            table.rows[7].cells[1].text = "Texnik"
            table.rows[7].cells[2].text = "444"
            table.rows[7].cells[3].text = "1"
            table.rows[7].cells[4].text = "-"
            doc.save(path)

            positions, warnings = parse_perechen_positions_6_4(path)
            perechen = parse_perechen(path)

        self.assertEqual([p["workplace_no"] for p in positions], [
            "000001", "000002", "000003", "000004",
        ])
        self.assertEqual(positions[0]["subdivision_headers"], ["PARENT", "CHILD"])
        self.assertEqual(positions[1]["group_index"], 0)
        self.assertEqual(positions[2]["subdivision_headers"], ["PARENT"])
        self.assertEqual(positions[2]["group_index"], 1)
        self.assertEqual(positions[3]["group_index"], 1)
        self.assertIn("000004", perechen)
        self.assertTrue(any("восстановлен" in warning for warning in warnings))

    def test_duplicate_card_number_is_repaired_only_by_unique_position_match(self):
        workplaces = [
            {"workplace_no": "016015", "position": "Rentgen laboranti", "source_file": "a.docx", "flags": []},
            {"workplace_no": "016015", "position": "2-toifali muhandis", "source_file": "b.docx", "flags": []},
        ]
        perechen = {
            "016015": {"position": "Rentgen laboranti"},
            "016016": {"position": "2-toifali muhandis"},
        }
        warnings: list[str] = []

        _repair_duplicate_workplace_numbers_from_perechen(workplaces, perechen, warnings)

        self.assertEqual([w["workplace_no"] for w in workplaces], ["016015", "016016"])
        self.assertIn("workplace_no_repaired_from_perechen", workplaces[1]["flags"])
        self.assertEqual(len(warnings), 1)


class Render64Tests(unittest.TestCase):
    def test_render_rebuilds_body_in_perechen_order_and_keeps_hierarchy(self):
        company = {
            "name": "SINOV KORXONASI",
            "parent": "BOSH TASHKILOT",
            "address": "Toshkent shahri",
            "product": "Sinov mahsuloti",
            "stir": "123456789",
            "ifut": "12345",
            "mxbt": "67890",
        }
        common = {
            "factors": {"overall": "2.0"},
            "injury_risk_class_6_4": "2",
            "ppe_status_6_4": "mos",
            "ppe_not_envisaged_6_4": False,
            "female_count_6_4": "0",
        }
        workplaces = [
            {
                **common,
                "workplace_no": "000001",
                "subdivision_6_4": "Child unit",
                "subdivision_headers_6_4": ["PARENT", "Child unit"],
                "subdivision_group_6_4": 0,
                "perechen_order_6_4": 0,
                "employees_count_6_4": "2",
            },
            {
                **common,
                "workplace_no": "000002",
                "subdivision_6_4": "PARENT",
                "subdivision_headers_6_4": ["PARENT"],
                "subdivision_group_6_4": 1,
                "perechen_order_6_4": 1,
                "employees_count_6_4": "1",
            },
        ]

        with tempfile.TemporaryDirectory() as tmp:
            output = render_6_4(company, workplaces, Path(tmp) / "6_4.docx")
            doc = Document(output)

        self.assertEqual(len(doc.tables), 3)
        self.assertEqual(len(doc.tables[1].rows), 7)
        body = doc.tables[2]
        self.assertEqual([body.rows[i].cells[0].text for i in (3, 4, 8)], [
            "PARENT", "Child unit", "PARENT",
        ])
        self.assertEqual(body.rows[5].cells[1].text, "1")
        self.assertEqual(body.rows[9].cells[1].text, "1")
        parent_run = next(run for run in body.rows[3].cells[0].paragraphs[0].runs if run.text)
        child_run = next(run for run in body.rows[4].cells[0].paragraphs[0].runs if run.text)
        self.assertTrue(parent_run.bold)
        self.assertFalse(parent_run.italic)
        self.assertTrue(child_run.bold)
        self.assertTrue(child_run.italic)
        self.assertEqual(body.rows[5].cells[0].paragraphs[0].alignment, WD_ALIGN_PARAGRAPH.CENTER)
        self.assertEqual(body.rows[5].cells[1].paragraphs[0].alignment, WD_ALIGN_PARAGRAPH.CENTER)

        value_cell = doc.tables[0].rows[0].cells[2]
        self.assertEqual(value_cell.paragraphs[0].alignment, WD_ALIGN_PARAGRAPH.CENTER)
        self.assertTrue(all(run.bold for run in value_cell.paragraphs[0].runs if run.text))
        code_cell = doc.tables[0].rows[4].cells[1]
        self.assertEqual(code_cell.paragraphs[0].alignment, WD_ALIGN_PARAGRAPH.CENTER)
        self.assertTrue(all(run.bold for run in code_cell.paragraphs[0].runs if run.text))
        body_children = list(doc._element.body.iterchildren())
        table_indices = [i for i, child in enumerate(body_children) if child.tag.endswith("}tbl")]
        self.assertEqual(body_children[table_indices[-1] + 1].tag.rsplit("}", 1)[-1], "sectPr")
        between = body_children[table_indices[-2] + 1]
        self.assertTrue(between.xpath("./w:pPr/w:pageBreakBefore"))
        all_text = "\n".join(
            cell.text for table in doc.tables for row in table.rows for cell in row.cells
        )
        self.assertIn("ЯТҲВ", all_text)
        self.assertNotIn("Ятҳв", all_text)


if __name__ == "__main__":
    unittest.main()
