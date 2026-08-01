from __future__ import annotations

import tempfile
import unittest
from pathlib import Path

from docx import Document

from attestation.services.extract import (
    Doc,
    _extract_ppe_bases_6_6,
    canonical_workplace_no,
    parse_perechen,
    split_workplace_no,
)
from attestation.services.normalize import to_cyrillic
from attestation.services.render import plan_rows_6_6, plan_sections_6_6, render_6_6


class PlanRows66Tests(unittest.TestCase):
    def test_only_harmful_groups_are_included_and_workplaces_are_sorted(self):
        workplaces = [
            {
                "workplace_no": "000002",
                "position": "Muhandis",
                "factors": {"chem": "3.2", "aerosols": "3.1", "noise": "2.0"},
            },
            {
                "workplace_no": "000001",
                "position": "Laborant",
                "factors": {
                    "noise": "3.3",
                    "vibration_general": "3.1",
                    "biological": "3.0",
                },
            },
        ]

        rows = plan_rows_6_6(workplaces, lang="lat")

        self.assertEqual(len(rows), 3)
        self.assertEqual(rows[0][0], "Karta\n№000001")
        self.assertEqual(rows[1][0], "")
        self.assertEqual(rows[2][0], "Karta\n№000002")
        self.assertIn("nafas olish", rows[2][1])
        self.assertNotIn("Muhandis", rows[2][0])

    def test_subdivision_rows_and_severity_posture_come_from_card(self):
        workplaces = [{
            "workplace_no": "000004",
            "position": "Bosh buxgalter",
            "subdivision": "Boshqaruv apparati",
            "factors": {"severity": "3.1", "intensity": "3.3"},
            "plan_6_6": {
                "severity_posture": (
                    "Vaqti-vaqti bilan noqulay, qat`iy holatda bo`lish ish vaqtining 50% gacha"
                )
            },
        }]

        sections = plan_sections_6_6(workplaces, lang="lat")

        self.assertEqual(sections[0][0], "Boshqaruv apparati")
        rows = sections[0][1]
        self.assertEqual(len(rows), 2)
        self.assertIn("50% gacha", rows[0][1])
        self.assertEqual(rows[0][0], "Karta\n№000004")
        self.assertEqual(rows[1][0], "")

    def test_no_harmful_factors_produces_explanatory_row(self):
        rows = plan_rows_6_6(
            [{"workplace_no": "1", "factors": {"chem": "2.0", "overall": "2.0"}}],
            lang="lat",
        )
        self.assertEqual(len(rows), 1)
        self.assertIn("aniqlanmadi", rows[0][1])

    def test_cyrillic_chemistry_ignores_calculated_total_without_harmful_agent(self):
        workplace = {
            "workplace_no": "000001",
            "factors": {"chem": "3.1"},
            "substances": [{"name": "Йиғиш омили", "cls": "3.1"}],
        }

        cyr_rows = plan_rows_6_6([workplace], lang="cyr")
        lat_rows = plan_rows_6_6([workplace], lang="lat")

        self.assertIn("аниқланмади", cyr_rows[0][1])
        self.assertIn("nafas olish", lat_rows[0][1])

    def test_cyrillic_tutuq_and_initial_e(self):
        self.assertEqual(to_cyrillic("ta'sir me'yor Elektromagnit aerozol"),
                         "таъсир меъёр Электромагнит аэрозол")


class Render66Tests(unittest.TestCase):
    def test_render_fills_requisites_and_removes_reference_client_data(self):
        company = {
            "name": "SINOV KORXONASI",
            "parent": "BOSH TASHKILOT",
            "address": "Toshkent shahri",
            "product": "Sinov mahsuloti",
            "stir": "123456789",
            "ifut": "12345",
            "mxbt": "67890",
        }
        workplaces = [
            {
                "workplace_no": "000001",
                "position": "Laborant",
                "factors": {"noise": "3.2", "overall": "3.2"},
            }
        ]

        with tempfile.TemporaryDirectory() as tmp:
            output = render_6_6(company, workplaces, Path(tmp) / "6_6.docx", lang="lat")
            doc = Document(output)

        text = "\n".join(p.text for p in doc.paragraphs)
        text += "\n" + "\n".join(
            cell.text for table in doc.tables for row in table.rows for cell in row.cells
        )
        self.assertIn("SINOV KORXONASI", text)
        self.assertIn("123456789", text)
        self.assertIn("Karta\n№000001", text)
        self.assertNotIn("Xidirov", text)
        self.assertNotIn("Aliqulov", text)
        self.assertNotIn("308425864", text)
        self.assertEqual(len(doc.tables), 3)
        self.assertEqual(len(doc.tables[-1].rows), 4)
        self.assertEqual(doc.tables[-1].rows[2].cells[0].text, "—")

    def test_cyrillic_execution_and_ppe_basis_match_reference_shape(self):
        company = {
            "name": "SINOV",
            "ppe_mandatory_6_6": "Тармоқ ЯТҲВ меъёри",
            "ppe_additional_6_6": "Жамоа шартномаси",
        }
        workplaces = [{
            "workplace_no": "000001",
            "subdivision": "Бошқариш аппарати",
            "factors": {"noise": "3.1"},
        }]
        with tempfile.TemporaryDirectory() as tmp:
            output = render_6_6(company, workplaces, Path(tmp) / "6_6.docx", lang="cyr")
            table = Document(output).tables[-1]

        self.assertEqual(table.rows[2].cells[0].text, "Бошқариш аппарати")
        self.assertEqual(table.rows[3].cells[0].text, "Карта\n№000001")
        self.assertEqual(table.rows[3].cells[4].text, "Қўлланма")
        self.assertIn("Тармоқ ЯТҲВ меъёри", table.rows[4].cells[0].text)


class Extract66Tests(unittest.TestCase):
    def test_latin_and_cyrillic_a_suffix_are_one_workplace(self):
        self.assertEqual(canonical_workplace_no(17, "a"), "000017а")
        self.assertEqual(canonical_workplace_no(17, "а"), "000017а")
        self.assertEqual(split_workplace_no("000017a"), (17, "а"))

    def test_ppe_bases_are_extracted_from_card_table(self):
        doc = Doc(
            tables=[[ 
                ["Мажбурий", "Тармоқ меъёри"],
                ["", "Изоҳ"],
                ["Қўшимча", "Жамоа шартномаси"],
            ]],
            paragraphs=[],
        )
        self.assertEqual(
            _extract_ppe_bases_6_6(doc),
            {
                "ppe_mandatory_6_6": "Тармоқ меъёри",
                "ppe_additional_6_6": "Жамоа шартномаси",
            },
        )

    def test_perechen_canonicalizes_latin_a_suffix(self):
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "perechen.docx"
            doc = Document()
            table = doc.add_table(rows=6, cols=5)
            for row_no in range(4):
                table.rows[row_no].cells[0].text = "header"
            table.rows[4].cells[0].text = "000017a"
            table.rows[4].cells[1].text = "Operator"
            table.rows[4].cells[2].text = "1234"
            table.rows[5].cells[0].text = "000018а"
            table.rows[5].cells[1].text = "Muhandis"
            table.rows[5].cells[2].text = "5678"
            doc.save(path)

            result = parse_perechen(path)

        self.assertEqual(set(result), {"000017а", "000018а"})
        self.assertNotIn("000017a", result)


if __name__ == "__main__":
    unittest.main()
