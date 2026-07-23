"""Движок извлечения данных из карт и «Перечня».

Архитектурный приём: сначала каждую карту и «Перечень» раскладываем в чистые
структуры (таблицы как сетки строк/ячеек + абзацы), затем по декларативным
якорям (см. ``mapping.py``) достаём поля. Результат — одна запись на рабочее
место + блок реквизитов компании. Оба шаблона (5_1б, 6_5) потом заполняются
из этого единого набора.

Зависит только от python-docx и стандартной библиотеки.
"""
from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path

from docx import Document

from . import mapping as M
from .normalize import (
    canon_yesno,
    clean_substance_name,
    fold,
    fold_contains,
    fold_contains_all,
    is_empty,
    max_class,
    normalize_class,
    normalize_number,
    normalize_spaces,
)

# --- Имя файла карты: число + ЛЮБОЙ разделитель + название -------------------
# Допускаем точку, подчёркивание, дефис, пробел: «24а. Должность.doc»,
# «3_Умумий_...doc», «000004.Bosh mexanik.doc». Имя — лишь запасной источник;
# номер и должность берём из содержимого документа (имя бывает ненадёжным).
_CARD_FILENAME_RE = re.compile(
    r"^\s*(\d+)\s*([а-яёА-ЯЁa-zA-Z]?)\s*[.,_\-\s]\s*(.+?)\s*\.(?:doc|docx)$", re.IGNORECASE
)


def canonical_workplace_no(number: int, suffix: str = "") -> str:
    """Каноничный номер рабочего места: 6 цифр + опц. суффикс-вариант."""
    return f"{number:06d}{suffix}"


def split_workplace_no(wp: str) -> tuple[int | None, str]:
    """Разбить номер РМ на (число, буквенный суффикс): «000012а» → (12, 'а')."""
    m = re.match(r"^0*(\d+)\s*([^\d\s]*)\s*$", wp or "")
    if not m:
        return None, ""
    return int(m.group(1)), m.group(2)


def workplace_sort_key(wp: str) -> tuple[int, str]:
    """Ключ сортировки: 000011 < 000011а < 000011б < 000012 (суффикс — любая буква)."""
    num, suffix = split_workplace_no(wp)
    return (num if num is not None else 10**9, suffix)


def parse_card_filename(basename: str) -> tuple[int | None, str, str]:
    """Разобрать имя файла карты → (число, суффикс-вариант, название).

    Запасной источник: основной — содержимое документа. Возвращает
    ``(None, "", "")`` если имя не похоже на карту.
    """
    m = _CARD_FILENAME_RE.match(basename)
    if not m:
        return None, "", ""
    return int(m.group(1)), m.group(2), normalize_spaces(m.group(3))


# ---------------------------------------------------------------------------
# Чтение docx в простые структуры
# ---------------------------------------------------------------------------
@dataclass
class Doc:
    """Документ как набор таблиц-сеток и абзацев."""

    tables: list[list[list[str]]]  # таблица → строки → ячейки (текст)
    paragraphs: list[str]


def read_docx(path: str | Path) -> Doc:
    document = Document(str(path))
    tables: list[list[list[str]]] = []
    for tbl in document.tables:
        grid: list[list[str]] = []
        ncols = len(tbl.columns)
        for row in tbl.rows:
            cells = row.cells
            grid.append([normalize_spaces(cells[c].text) for c in range(min(ncols, len(cells)))])
        tables.append(grid)
    paragraphs = [normalize_spaces(p.text) for p in document.paragraphs]
    return Doc(tables=tables, paragraphs=paragraphs)


# --- Помощники поиска по якорям (двуязычные, через fold) ---------------------
def _contains(text: str, anchor: str) -> bool:
    return fold_contains(text, anchor)


def find_table(tables: list[list[list[str]]], anchor: str) -> list[list[str]] | None:
    """Первая таблица, в любой ячейке которой встречается anchor."""
    for grid in tables:
        for row in grid:
            if any(_contains(c, anchor) for c in row):
                return grid
    return None


def find_row(grid: list[list[str]], anchor: str) -> list[str] | None:
    for row in grid:
        if any(_contains(c, anchor) for c in row):
            return row
    return None


def find_row_tokens(grid: list[list[str]], tokens: list[str]) -> list[str] | None:
    """Первая строка, в тексте которой присутствуют ВСЕ токены (после fold)."""
    for row in grid:
        if fold_contains_all(" ".join(row), tokens):
            return row
    return None


def row_value_after_label(row: list[str], label: str) -> str:
    """Значение в строке «метка | значение»: первая непустая ячейка ≠ метки."""
    for cell in row:
        if not is_empty(cell) and not _contains(cell, label):
            return cell
    return ""


def last_value(row: list[str]) -> str:
    """Последняя непустая ячейка строки (аналог col -1)."""
    for cell in reversed(row):
        if not is_empty(cell):
            return cell
    return ""


# --- Сопоставление номеров разделов -----------------------------------------
_SECTION_RE = re.compile(r"^\s*(\d+(?:\.\d+)*)")


def _section_of(cell0: str) -> str | None:
    """Если ячейка col0 начинается с номера раздела (1.1, 1.3.2 …) — вернуть его."""
    m = _SECTION_RE.match(cell0)
    return m.group(1) if m else None


def section_matches(section: str | None, prefix: str) -> bool:
    """Раздел принадлежит префиксу? (1.1.3 ∈ 1.1, но 1.11 ∉ 1.1)."""
    if not section:
        return False
    return section == prefix or section.startswith(prefix + ".")


# ---------------------------------------------------------------------------
# Реквизиты компании
# ---------------------------------------------------------------------------
def _reqs_value(row: list[str], anchor: str) -> str:
    """Значение реквизита: первая непустая ячейка ≠ метки. Сохраняет «-»."""
    flabel = fold(anchor)
    for cell in row:
        raw = normalize_spaces(cell)
        if not raw:
            continue
        if flabel and flabel in fold(cell):  # это ячейка-метка
            continue
        return raw
    return ""


# Метки кодов в двух письмах; берём цифры (даже с пробелами) после метки
_CODE_LABELS = (
    ("stir", r"(?:СТИР|STIR)"),
    ("ifut", r"(?:ИФУТ|IFUT)"),
    ("mxbt", r"(?:МХБТ|MX[БB]T)"),
)


def _extract_codes(grid: list[list[str]]) -> dict:
    text = normalize_spaces(" ".join(" ".join(row) for row in grid))
    out = {}
    for key, lab in _CODE_LABELS:
        m = re.search(lab + r"\s*[:\-]?\s*([0-9][0-9 ]*)", text, re.IGNORECASE)
        out[key] = re.sub(r"\s+", "", m.group(1)) if m else ""
    return out


def extract_company_data(doc: Doc) -> dict:
    grid = find_table(doc.tables, M.COMPANY_FIELDS["name"])
    data = {k: "" for k in (*M.COMPANY_FIELDS, "stir", "ifut", "mxbt")}
    if grid is None:
        return data

    for key, anchor in M.COMPANY_FIELDS.items():
        row = find_row(grid, anchor)
        if row:
            data[key] = _reqs_value(row, anchor)

    data.update(_extract_codes(grid))
    return data


# ---------------------------------------------------------------------------
# Разбор большой таблицы факторов
# ---------------------------------------------------------------------------
@dataclass
class FactorRow:
    section: str | None  # номер раздела, «перенесённый» сверху
    c0: str              # исходный col0 (пусто у под-строк)
    name: str            # col1 — фактор/вещество
    actual: str          # «Ҳақиқий даражаси»
    duration: str        # «Таъсир этиш давомийлиги (соат/%)»
    cls: str             # класс (последняя колонка)
    norm: str = ""       # «Гигиеник меъёр (РЭЧК, РЭЧД)» — норма (для Excel-протоколов)


def _find_factor_table(doc: Doc) -> list[list[str]] | None:
    """Большая таблица факторов: по якорю заголовка, иначе самая длинная."""
    for grid in doc.tables:
        for row in grid[:3]:
            joined = " ".join(row).lower()
            if all(a in joined for a in (M.FACTOR_HEADER_CLASS,)) and any(
                a in joined for a in M.FACTOR_TABLE_HEADER_ANCHORS
            ):
                return grid
    return max(doc.tables, key=len) if doc.tables else None


def _factor_columns(grid: list[list[str]]) -> tuple[int, int, int, int, int]:
    """Индексы колонок (name, actual, duration, class, norm) по строке-заголовку.

    ``norm`` (гигиеник меъёр) добавлен для Excel-протоколов лаб. замеров; в
    подтверждённой раскладке карты это колонка 2 (name=1, norm=2, actual=3,
    duration=4, class=последняя). Возвращается ПОСЛЕДНИМ, чтобы не сдвигать
    существующие индексы (в т.ч. class_col=[3]).
    """
    name_col, norm_col, actual_col, duration_col = 1, 2, 3, 4  # дефолты по раскладке
    class_col = max(len(r) for r in grid) - 1
    norm_found = False
    for row in grid[:3]:
        for i, c in enumerate(row):
            if _contains(c, M.FACTOR_HEADER_ACTUAL):
                actual_col = i
            if _contains(c, M.FACTOR_HEADER_DURATION):
                duration_col = i
            if _contains(c, "омиллари"):
                name_col = i
            # Норма («меъёр») — берём ПЕРВУЮ (левую) такую колонку и фиксируем.
            if not norm_found and _contains(c, M.FACTOR_HEADER_NORM):
                norm_col = i
                norm_found = True
        if any(_contains(c, M.FACTOR_HEADER_CLASS) for c in row):
            class_col = len(row) - 1
    return name_col, actual_col, duration_col, class_col, norm_col


def _parse_factor_rows(grid: list[list[str]]) -> list[FactorRow]:
    name_col, actual_col, duration_col, class_col, norm_col = _factor_columns(grid)
    rows: list[FactorRow] = []
    current_section: str | None = None
    for raw in grid:
        if not raw:
            continue
        c0 = raw[0] if len(raw) > 0 else ""
        sec = _section_of(c0)
        if sec:
            current_section = sec

        def cell(i: int) -> str:
            return raw[i] if i < len(raw) else ""

        rows.append(
            FactorRow(
                section=current_section,
                c0=c0,
                name=cell(name_col),
                actual=cell(actual_col),
                duration=cell(duration_col),
                cls=last_value(raw) if class_col >= len(raw) else cell(class_col),
                norm=cell(norm_col),
            )
        )
    return rows


def _factor_values(grid: list[list[str]], rows: list[FactorRow]) -> dict:
    """Классы по всем факторным колонкам 6_5."""
    factors: dict[str, str] = {}
    class_col = _factor_columns(grid)[3]

    # Подытоги/общий класс — по набору токенов, значение в КОЛОНКЕ КЛАССА.
    # Берём именно класс-колонку (а не последнюю непустую): если класс пуст,
    # последней непустой была бы объединённая ячейка-метка («Zararli
    # moddalarni umumiy baholash») — тогда вместо «-» протекал бы текст.
    for key, tokens in M.FACTOR_SUBTOTALS.items():
        row = find_row_tokens(grid, tokens)
        if not row:
            factors[key] = "-"
        else:
            val = row[class_col] if class_col < len(row) else last_value(row)
            factors[key] = normalize_class(val)

    # Факторы без подытога — максимум класса среди строк раздела
    for key, prefixes in M.FACTOR_SECTIONS.items():
        classes = [
            fr.cls
            for fr in rows
            if any(section_matches(fr.section, p) for p in prefixes)
            and not is_empty(fr.cls)
        ]
        factors[key] = max_class(classes)
    return factors


def _extract_substances(rows: list[FactorRow]) -> list[dict]:
    """Вредные вещества для 5_1б: под-строки разделов 1.1.* с реальным значением.

    Возвращаются в каноничном порядке эталона (см. ``SUBSTANCE_ORDER``):
    углерод оксиди первым, затем по приоритету; неизвестные — в порядке карты.
    """
    substances: list[dict] = []
    for fr in rows:
        if not section_matches(fr.section, M.SUBSTANCE_SECTION_PREFIX):
            continue
        # Вещество — под-строка (col0 пуст) с измеренным «Ҳақиқий даражаси»
        if is_empty(fr.c0) and not is_empty(fr.actual) and not is_empty(fr.name):
            name = clean_substance_name(fr.name)
            pct = normalize_number(fr.duration)
            if name:
                # name/pct — для 5_1б (не менять); norma/actual/cls — для Excel-протокола 1.
                substances.append({
                    "name": name,
                    "pct": pct,
                    "norma": normalize_number(fr.norm),
                    "actual": normalize_number(fr.actual),
                    "cls": normalize_class(fr.cls),
                })

    def order_key(item: tuple[int, dict]) -> tuple[int, int]:
        idx, s = item
        rank = M.SUBSTANCE_ORDER.get(s["name"].lower(), M.SUBSTANCE_ORDER_UNKNOWN)
        return (rank, idx)

    return [s for _, s in sorted(enumerate(substances), key=order_key)]


# ---------------------------------------------------------------------------
# Построчные (сырые) замеры для Excel-протоколов лабораторных замеров
# ---------------------------------------------------------------------------
# В отличие от factors (подытоги/max-класс по разделам) — здесь берётся ИМЕННО
# строка нужного раздела карты с её нормой/фактом/временем/классом. Каждый замер:
#   {"norma", "actual", "time", "cls"}  (строки; норма может быть диапазоном
#    «23,0-31,0», факт — числом или текстом «йўқ»; «-» = замера нет).
def _measurement(fr: FactorRow) -> dict:
    """Замер (норма/факт/время/класс) из строки факторной таблицы."""
    return {
        "norma": normalize_number(fr.norm),
        "actual": normalize_number(fr.actual),
        "time": normalize_number(fr.duration),
        "cls": normalize_class(fr.cls),
    }


def _row_by_section(
    rows: list[FactorRow], section: str, *,
    prefix: bool = False, need_actual: bool = False, name_anchor: str | None = None,
) -> FactorRow | None:
    """Найти строку раздела: exact/prefix-номер, опц. с непустым фактом и якорем имени."""
    for fr in rows:
        ok = section_matches(fr.section, section) if prefix else (fr.section == section)
        if not ok:
            continue
        if name_anchor and not fold_contains(fr.name, name_anchor):
            continue
        if need_actual and is_empty(fr.actual):
            continue
        return fr
    return None


def _measurement_if_actual(fr: FactorRow | None) -> dict | None:
    """Замер, если факт непуст (иначе None — под-строка остаётся пустой рамкой)."""
    if fr is None or is_empty(fr.actual):
        return None
    return _measurement(fr)


def _extract_physical_measurements(rows: list[FactorRow]) -> dict:
    """Файл 2: шум 1.3.2, вибрация локальная/общая 1.3.5/1.3.6, инфразвук 1.3.1."""
    # Шум: строка 1.3.2 именно «Shovqin tovush darajasi… dBA» (в карте под номером
    # 1.3.2 идёт ещё строка инфразвука — её не берём). Фолбэк — любая 1.3.2 с фактом.
    noise = _row_by_section(rows, M.PHYSICAL_NOISE_SECTION,
                            name_anchor=M.PHYSICAL_NOISE_NAME_ANCHOR)
    if noise is None:
        noise = _row_by_section(rows, M.PHYSICAL_NOISE_SECTION, need_actual=True)
    out = {"noise": _measurement_if_actual(noise)}
    for key, sec in M.PHYSICAL_SECTIONS.items():
        out[key] = _measurement_if_actual(
            _row_by_section(rows, sec, prefix=True, need_actual=True)
        )
    return out


def _strip_category_unit(name: str) -> str:
    """«Ib – 88(78-97), Vt/m» → «Ib – 88(78-97)» (срез единицы в метке категории)."""
    s = normalize_spaces(name)
    low = s.lower()
    for marker in M.MICROCLIMATE_CATEGORY_UNIT_MARKERS:
        pos = low.find(marker)
        if pos != -1:
            return s[:pos].strip()
    return s


def _extract_microclimate_measurements(rows: list[FactorRow]) -> dict:
    """Файл 3: активная категория 1.8.x + температура/скорость/влажность/теплоизлучение.

    Активна та строка 1.8.1..1.8.5, у которой заполнен факт. Метка «Ishlar toifasi»
    («Iб – 88(78-97)») берётся из параллельного раздела 1.7 (WBGT) с тем же номером.
    """
    temp = None
    category_label = ""
    for idx, sec in enumerate(M.MICROCLIMATE_TEMP_SECTIONS):
        fr = _row_by_section(rows, sec, need_actual=True)
        if fr is not None:
            temp = _measurement(fr)
            cat_fr = _row_by_section(rows, M.MICROCLIMATE_CATEGORY_SECTIONS[idx])
            if cat_fr and not is_empty(cat_fr.name):
                category_label = _strip_category_unit(cat_fr.name)
            break
    return {
        "category_label": category_label,
        "temp": temp,
        "air_speed": _measurement_if_actual(
            _row_by_section(rows, M.MICROCLIMATE_AIR_SPEED_SECTION)),
        "humidity": _measurement_if_actual(
            _row_by_section(rows, M.MICROCLIMATE_HUMIDITY_SECTION)),
        # Теплоизлучение: факт может быть текстом «йўқ» (не число) — сохраняем как есть.
        "heat_radiation": _measurement_if_actual(
            _row_by_section(rows, M.MICROCLIMATE_HEAT_SECTION)),
    }


_LIGHTING_DISCHARGE_RE = re.compile(r"\b(VIII|VII|III|VI|IV|IX|II|V|I)\b\s*$")


def _extract_lighting_measurements(rows: list[FactorRow]) -> dict:
    """Файл 4: разряд зрит. работ + КЕО (естеств./смеш.), искусств. лк, пульсация."""
    discharge = ""
    disc_row = next(
        (fr for fr in rows if fold_contains(fr.name, M.LIGHTING_DISCHARGE_ANCHOR)), None
    )
    if disc_row:
        m = _LIGHTING_DISCHARGE_RE.search(normalize_spaces(disc_row.name))
        if m:
            discharge = m.group(1)
    out = {"discharge": discharge}
    for key, sec in M.LIGHTING_SECTIONS.items():
        out[key] = _measurement_if_actual(_row_by_section(rows, sec))
    return out


def _extract_em_measurements(rows: list[FactorRow]) -> dict:
    """Файл 5: замеры ЭМИ 1.4.1..1.4.10 (у медиков заполнена обычно только 1.4.10)."""
    out: dict[str, dict | None] = {}
    for sec in M.EM_SECTIONS:
        out[sec] = _measurement_if_actual(_row_by_section(rows, sec))
    return out


# ---------------------------------------------------------------------------
# Подразделение, должность, номер
# ---------------------------------------------------------------------------
def _extract_position_from_doc(doc: Doc) -> str:
    """Должность — следующий непустой абзац после якоря «ИШ ЖОЙИ НОМИ».

    Работает и для кириллицы (``…ИШ ЖОЙИ НОМИ`` / ``Бош бухгалтер``), и для
    латиницы (``…ISH JOYI NOMI`` / ``Stropalchi``) — якорь матчится через fold.
    """
    paras = doc.paragraphs
    for i, p in enumerate(paras):
        if _contains(p, M.POSITION_ANCHOR):
            for q in paras[i + 1:]:
                if not is_empty(q):
                    return q
    return ""


# Номер РМ из содержимого: «000009а-сонли» / «000012-сонли» / «000001-sonli».
# Берём по СЫРОМУ тексту (не fold), чтобы сохранить кириллический суффикс «а».
_WORKPLACE_CONTENT_RE = re.compile(
    r"([0-9]{4,6})\s*([а-яёА-ЯЁa-zA-Z]?)\s*[-—–]?\s*(?:сонли|sonli)", re.IGNORECASE
)


def _extract_workplace_no_from_doc(doc: Doc) -> tuple[str, str]:
    """Вернуть (число, суффикс-вариант) из строки «…-сонли» документа."""
    for p in doc.paragraphs:
        m = _WORKPLACE_CONTENT_RE.search(p)
        if m:
            return m.group(1), m.group(2).lower()
    for grid in doc.tables:
        for row in grid:
            m = _WORKPLACE_CONTENT_RE.search(" ".join(row))
            if m:
                return m.group(1), m.group(2).lower()
    return "", ""


def canonical_subdivision(name: str) -> str:
    """Привести название подразделения к каноничному (см. ``SUBDIVISION_ALIASES``).

    Нужно, чтобы один физический отдел, по-разному названный в разных картах
    («Рентген кабинет» / «Рентгент бўлими»), в сводной 6_4 стал ОДНИМ блоком.
    Сопоставление через ``fold`` — регистр/письмо варианта не важны.
    """
    f = fold(name)
    for variant, cyr in M.SUBDIVISION_ALIASES.items():
        if fold(variant) == f:
            return cyr
    return name


def _extract_subdivision(doc: Doc) -> str:
    grid = find_table(doc.tables, M.SUBDIVISION_TABLE_ANCHOR)
    if grid is None:
        return ""
    row = find_row(grid, M.SUBDIVISION_ROW_ANCHOR)
    raw = row_value_after_label(row, M.SUBDIVISION_ROW_ANCHOR) if row else ""
    return canonical_subdivision(raw)


def _extract_employee_counts(doc: Doc) -> tuple[str, str]:
    """(ишловчилар сони, шу жумладан аёллар) из таблицы «Таркибий бўлинма».

    Нужно для 6_4 (сводная қайднома по подразделениям): там строки не только
    «сколько рабочих мест», но и «сколько на них занято людей» / «из них
    женщин». Строка гендерной разбивки есть не во всех вариантах карты —
    тогда возвращаем "" (НЕ "0"), чтобы при агрегации в render_6_4 отличить
    «неизвестно» от «действительно ноль».

    В основном варианте карты «Ишловчилар сони» и «Аёллар» — ОДНА строка
    («Ишловчилар сони | 30 | | - | | 30»), а подписи колонок («Иш жойида» /
    «Барча аналогик иш жойларида» / «Аёллар») — СТРОКОЙ НИЖЕ. Отдельного
    якоря «жумладан аёллар» в этой строке нет, поэтому берём колонку под
    подписью «Аёллар» из строки заголовков колонок. «-» в этой колонке —
    «женщин нет» (настоящий ноль), а не «нет данных», поэтому НЕ используем
    ``last_value`` (он трактует «-» как пусто и пропускает такую колонку).
    """
    grid = find_table(doc.tables, M.SUBDIVISION_TABLE_ANCHOR)
    if grid is None:
        return "", ""
    workers_row = find_row(grid, M.WORKERS_ROW_ANCHOR)
    workers = normalize_number(row_value_after_label(workers_row, M.WORKERS_ROW_ANCHOR)) if workers_row else ""

    female_row = find_row(grid, M.WORKERS_FEMALE_ROW_ANCHOR)
    if female_row:
        female = normalize_number(row_value_after_label(female_row, M.WORKERS_FEMALE_ROW_ANCHOR))
    else:
        female = ""
        if workers_row:
            idx = grid.index(workers_row)
            header_row = grid[idx + 1] if idx + 1 < len(grid) else None
            if header_row:
                for i, cell in enumerate(header_row):
                    if _contains(cell, M.WORKERS_FEMALE_COL_ANCHOR) and i < len(workers_row):
                        female = normalize_number(workers_row[i])
                        break
    return workers, female


# ---------------------------------------------------------------------------
# СИЗ (ЯТҲВ) и льготы
# ---------------------------------------------------------------------------
def _extract_ppe(doc: Doc) -> str:
    """ЯТҲВ обеспеченность: «ҳа», если хоть где-то отмечено наличие (бор/ҳа).

    Внимание: в шапке таблицы 6 подстрока «мавжудлиги» встречается дважды —
    в нужной колонке «Ходимда ЯТҲВ мавжудлиги (йўқ/бор)» и в последней
    «Сертификат … мавжудлиги». Берём ПЕРВОЕ совпадение (нужная колонка левее).
    """
    grid = find_table(doc.tables, M.PPE_TABLE_ANCHOR)
    if grid is None:
        return "йўқ"
    value_col = None
    for row in grid[:2]:
        for i, c in enumerate(row):
            if _contains(c, M.PPE_ROW_VALUE_ANCHOR):
                value_col = i
                break
        if value_col is not None:
            break
    has_ppe = False
    for row in grid[1:]:  # пропускаем шапку
        candidates = [row[value_col]] if value_col is not None and value_col < len(row) else row
        for c in candidates:
            if not is_empty(c) and canon_yesno(c) == "ҳа":
                has_ppe = True
    return "ҳа" if has_ppe else "йўқ"


# --- 6_4: реальные итоговые показатели (см. mapping.INJURY_RISK_CLASS_ANCHORS
# / PPE_ANSWER_ANCHOR) — НЕ путать с injury_risk/ppe_provided выше: те эвристика
# и таблица СИЗ соответственно, используются только 5_1б/6_5/экраном ревью.
_INJURY_CLASS_VALUE_RE = re.compile(r"[_\s]([1-3])[_\s]*(?:дараж|daraja)", re.IGNORECASE)
# Заполненный пропуск «____2____» — единственная часть п.2.3, одинаковая у ВСЕХ
# известных формулировок (см. mapping.INJURY_RISK_KEYWORDS), поэтому основной путь.
_INJURY_CLASS_BLANK_RE = re.compile(r"_{2,}\s*([1-3])\s*_{2,}")
# Номер пункта 2.3 в начале абзаца — отделяет итоговый класс от таблицы п.2.1,
# где тоже встречаются одиночные цифры 1/2/3 (оценки по отдельным факторам).
_SECTION_2_3_RE = re.compile(r"^\s*2\.3[.\s]")


def _extract_injury_risk_class_6_4(doc: Doc) -> str:
    """Класс травмоопасности (1-3) из п.2.3 карты; "" если не найден.

    Пункт опознаём по номеру «2.3» + любому из ключевых слов травмоопасности
    (``INJURY_RISK_KEYWORDS``) — формулировка различается у клиентов:
        «…шикастлаш хавфлилиги бўйича ____2____даражага мансуб»
        «…Шикастланиш хавфи бўйича … синфга тегишли: ____2____»
    В обоих случаях значение — вписанное в пропуск число «____N____», поэтому
    сначала ищем именно пропуск, и лишь затем (запасной путь, для карт без
    подчёркиваний) — число перед словом «даража/daraja».

    Специально НЕ используем дефолт при промахе: пустая строка означает
    «не извлечено», и render_6_4 громко предупредит об этом, а не тихо отнесёт
    рабочее место к произвольному классу.
    """
    for p in doc.paragraphs:
        is_section = bool(_SECTION_2_3_RE.match(p)) and any(
            _contains(p, kw) for kw in M.INJURY_RISK_KEYWORDS
        )
        # Запасное опознание для карт, где номер пункта потерялся при конвертации
        legacy = _contains(p, M.INJURY_RISK_CLASS_ANCHORS[0]) and _contains(
            p, M.INJURY_RISK_CLASS_ANCHORS[1]
        )
        if not (is_section or legacy):
            continue
        m = _INJURY_CLASS_BLANK_RE.search(p) or _INJURY_CLASS_VALUE_RE.search(p)
        if m:
            return m.group(1)
    return ""


_FOLD_YES_TOKENS = {"ha", "bor", "xa", "da", "yes"}
_FOLD_NO_TOKENS = {"yoq", "yuq", "yok", "no"}


def _fold_yesno_token(text: str) -> str:
    """«ha»/«йўқ»-подобный токен → 'ha'/'yoq'; иначе "" (не распознано)."""
    f = fold(text)
    if f in _FOLD_YES_TOKENS:
        return "ha"
    if f in _FOLD_NO_TOKENS:
        return "yoq"
    return ""


def _extract_ppe_status_6_4(doc: Doc) -> str:
    """Статус ЯТҲВ из п.3.3 карты: 'mos' / 'mos_emas' / 'kutilmagan'.

    П.3.3 в этом шаблоне карты — УТВЕРЖДЕНИЕ: «Иш жойи ЯТҲВ билан
    таъминланганлик талабларига жавоб беради» («…ОТВЕЧАЕТ требованиям»). Это
    статичный текст без переменной ha/yo'q — сам факт его наличия означает
    «Мос» (соответствует). Проверено на эталоне клиента: колонка «Мос» = 100%
    рабочих мест по всем подразделениям, «Мос эмас» = 0. Прежняя трактовка
    «нет ответа → Ятҳв кўзда тутилмаган» была ошибкой (утверждение читалось как
    незаполненное поле); ещё более ранняя ошибка — брать ha/yo'q из раздела IV
    (льготы/пенсия), к обеспеченности СИЗ отношения не имеющего.

    Если в конкретной карте после текста п.3.3 всё же вписан явный ответ
    (инлайн или в одном из следующих 2 абзацев, до раздела IV «Кафолатлар»),
    он имеет приоритет: ha → 'mos', yo'q → 'mos_emas'. 'kutilmagan' остаётся
    лишь для карт, где раздела 3.3 вовсе нет (обеспеченность СИЗ не оценивалась).
    """
    paras = doc.paragraphs
    for i, p in enumerate(paras):
        if not _contains(p, M.PPE_ANSWER_ANCHOR):
            continue
        fp, fa = fold(p), fold(M.PPE_ANSWER_ANCHOR)
        tail = fp.split(fa, 1)[1] if fa in fp else ""
        answer = _fold_yesno_token(tail)
        if not answer:
            for cand in paras[i + 1:i + 3]:
                if is_empty(cand):
                    continue
                if _contains(cand, M.PPE_NEXT_SECTION_STOP_ANCHOR):
                    break
                answer = _fold_yesno_token(cand)
                break
        if answer == "yoq":
            return "mos_emas"
        # Явный «ha» ИЛИ статичное утверждение п.3.3 без переменной → «Мос».
        return "mos"
    return "kutilmagan"


def _extract_ppe_not_envisaged_6_4(doc: Doc) -> bool:
    """«Ятҳв кўзда тутилмаган» (СИЗ не предусмотрены) — из таблицы СИЗ (п.3.2).

    В таблице СИЗ наименование каждого средства — во 2-й колонке. Если для
    рабочего места СИЗ не предусмотрены, вместо названий во всех строках стоит
    literal «Кўзда тутилмаган», т.е. НЕТ ни одного реального наименования СИЗ.
    Тогда место идёт в колонку c11 сводной 6_4 «Ятҳв кўзда тутилмаган».

    Это НЕЗАВИСИМАЯ от «Мос/Мос эмас» (п.3.3) величина: рабочее место может быть
    одновременно «Мос» (формально соответствует требованиям) и «кўзда тутилмаган»
    (СИЗ не предусмотрены). Сверено с эталоном клиента: колонка «кўзда» = 93 по
    компании при «Мос» = 114 (все) — суммы намеренно пересекаются.
    """
    grid = find_table(doc.tables, M.PPE_TABLE_ANCHOR)
    if grid is None:
        return False
    for row in grid[1:]:  # пропускаем шапку
        name = normalize_spaces(row[1]) if len(row) > 1 else ""
        if not name or name.endswith(":"):
            continue  # категория «Мажбурий:» / «Қўшимча:», не наименование СИЗ
        if fold_contains(name, M.PPE_NOT_ENVISAGED_MARKER):
            continue  # «Кўзда тутилмаган» — СИЗ нет
        return False  # найдено реальное наименование СИЗ → предусмотрены
    return True  # реальных СИЗ в таблице нет → «кўзда тутилмаган»


def _extract_benefits(doc: Doc) -> dict:
    grid = find_table(doc.tables, M.BENEFITS_TABLE_ANCHOR)
    result = {k: "йўқ" for k in M.BENEFITS_ROWS}
    if grid is None:
        return result
    # Колонка фактического наличия: «Ҳақиқатда мавжудлиги» ИЛИ «Ҳақиқий
    # мавжудлиги» (формулировка различается у клиентов) — берём ПЕРВУЮ колонку
    # с «мавжудлиги» (она левее «ўрнатиш зарурлиги»/«асос»).
    fact_col = None
    for row in grid[:2]:
        for i, c in enumerate(row):
            if _contains(c, M.BENEFITS_FACT_COL_ANCHOR):
                fact_col = i
                break
        if fact_col is not None:
            break
    for key, anchor in M.BENEFITS_ROWS.items():
        for row in grid:
            if any(_contains(c, anchor) for c in row):
                val = row[fact_col] if fact_col is not None and fact_col < len(row) else ""
                result[key] = canon_yesno(val)
                break
    return result


# ---------------------------------------------------------------------------
# Эвристики для полей без детерминированного источника
# ---------------------------------------------------------------------------
def is_medical(position: str, subdivision: str) -> bool:
    """Медицинское/врачебное рабочее место?

    Опираемся и на должность, и на подразделение: группа
    «Врачлар, ўрта ва кичик тиббиёт ходимлари» содержит «тиббиёт»/«врач»,
    что относит к медицине даже немедицинские должности внутри неё
    (напр. «Хўжалик бекаси»), а директор-«Шифокор» ловится по должности.
    """
    text = f"{position} {subdivision}".lower()
    return any(kw in text for kw in M.MEDICAL_KEYWORDS)


def _injury_risk(medical: bool) -> str:
    return M.INJURY_RISK_MEDICAL if medical else M.INJURY_RISK_DEFAULT


def _extract_pension(doc: Doc) -> str:
    """Льготная пенсия — из раздела 4.2 карты: «…пенсия таъминоти ҳуқуқи __ ҳа __».

    Значение вписано в пропуск между «ҳуқуқи»/«huquqi» и скобкой «(агар…)».
    """
    for p in doc.paragraphs:
        if not (_contains(p, "пенсия таъминоти") or _contains(p, "пенсия таминоти")):
            continue
        m = re.search(r"(?:ҳуқуқи|хуқуқи|huquqi)(.*?)(?:\(|$)", p, re.IGNORECASE)
        if m:
            val = re.sub(r"[_‎\s]+", " ", m.group(1)).strip()
            return canon_yesno(val)
    return "йўқ"


# ---------------------------------------------------------------------------
# Сборка записи по одной карте
# ---------------------------------------------------------------------------
def extract_card(docx_path: str | Path, basename: str) -> dict:
    """Извлечь запись рабочего места из одной карты (.docx)."""
    doc = read_docx(docx_path)

    # Номер и должность — ИЗ СОДЕРЖИМОГО документа (имя файла ненадёжно:
    # бывает через подчёркивание «3_…», или вовсе не совпадает — «Energetik»
    # в имени при «Montajchi» внутри). Имя файла — лишь запасной источник.
    _f_num, f_suffix, f_title = parse_card_filename(basename)
    doc_num, doc_suffix = _extract_workplace_no_from_doc(doc)

    # Номер — ТОЛЬКО из содержимого: имя файла ненадёжно (подчёркивания, чужое
    # название). Нет строки «…-сонли» → номер пуст → карта уйдёт в «НЕ РАСПОЗНАНЫ»
    # (громкое предупреждение в pipeline), а не получит случайный номер из имени.
    suffix = doc_suffix or f_suffix  # суффикс-вариант обычно есть и в содержимом
    workplace_no = canonical_workplace_no(int(doc_num), suffix) if doc_num else ""

    position = _extract_position_from_doc(doc) or f_title
    subdivision = _extract_subdivision(doc)

    grid = _find_factor_table(doc)
    factor_rows = _parse_factor_rows(grid) if grid else []
    factors = _factor_values(grid, factor_rows) if grid else {}
    substances = _extract_substances(factor_rows)

    # Построчные замеры для Excel-протоколов лаб. замеров (файлы 2–5). Файл 1
    # берёт norma/actual/cls из substances выше. См. mapping.*_SECTIONS.
    physical_measurements = _extract_physical_measurements(factor_rows)
    microclimate_measurements = _extract_microclimate_measurements(factor_rows)
    lighting_measurements = _extract_lighting_measurements(factor_rows)
    em_measurements = _extract_em_measurements(factor_rows)

    # Льготы (колонки 21–25 в 6_5) — ПРЯМО ИЗ КАРТЫ, раздел 4 (НЕ эвристики):
    #   отпуск/сокр.день/питание/молоко — таблица 4.1, колонка «Ҳақиқатда мавжудлиги»;
    #   льготная пенсия — текст 4.2 («…пенсия таъминоти ҳуқуқи __ ҳа __»).
    # Прим.: эталон Бухоро вручную ставит молоко=ҳа всем медикам (20 РМ) вопреки
    # форме карты — эти ячейки в самопроверке Бухоро идут в «заметки». См. README.
    benefits = _extract_benefits(doc)
    ppe = _extract_ppe(doc)
    medical = is_medical(position, subdivision)
    injury_risk = _injury_risk(medical)
    pension = _extract_pension(doc)
    employees_count, female_count = _extract_employee_counts(doc)

    # 6_4-специфичные показатели (см. Шаг 4 спецификации) — читаются НАПРЯМУЮ
    # из карты, независимо от injury_risk/ppe_provided выше (те эвристика/
    # таблица СИЗ для 5_1б/6_5, менять их поведение не нужно).
    injury_risk_class_6_4 = _extract_injury_risk_class_6_4(doc)
    ppe_status_6_4 = _extract_ppe_status_6_4(doc)
    ppe_not_envisaged_6_4 = _extract_ppe_not_envisaged_6_4(doc)

    # Флаги для подсветки в UI (поля, требующие проверки оператором)
    flags: list[str] = []
    if not factors.get("overall") or factors.get("overall") == "-":
        flags.append("overall_missing")
    flags.append("injury_risk_heuristic")     # травмоопасность (c18) — эвристика
    if not employees_count:
        # Нет строки «ишловчилар сони» в карте — считаем минимум 1 (карта
        # описывает занятое рабочее место), но подсвечиваем для проверки.
        employees_count = "1"
        flags.append("employees_count_missing")
    if not female_count:
        flags.append("female_count_missing")  # нужно для 6_4; см. mapping.WORKERS_FEMALE_ROW_ANCHOR

    return {
        "workplace_no": workplace_no,
        "source_file": basename,
        "workplace_no_in_doc": (doc_num + doc_suffix) if doc_num else "",
        "position": position,
        "position_from_filename": f_title,
        "subdivision": subdivision,
        "job_code": "",  # заполняется из «Перечня» на этапе слияния
        "factors": factors,
        "substances": substances,
        "ppe_provided": ppe,
        "benefits": benefits,
        "injury_risk": injury_risk,
        "privileged_pension": pension,
        "employees_count": employees_count,
        "female_count": female_count,
        "injury_risk_class_6_4": injury_risk_class_6_4,
        "ppe_status_6_4": ppe_status_6_4,
        "ppe_not_envisaged_6_4": ppe_not_envisaged_6_4,
        # Построчные замеры для Excel-протоколов (файлы 2–5); файл 1 — из substances.
        "physical_measurements": physical_measurements,
        "microclimate_measurements": microclimate_measurements,
        "lighting_measurements": lighting_measurements,
        "em_measurements": em_measurements,
        "flags": flags,
    }


# ---------------------------------------------------------------------------
# Разбор «Перечня» — источник кода должности (и сверка должности)
# ---------------------------------------------------------------------------
# Номер рабочего места в «Перечне»: 4-6 цифр + опц. буквенный суффикс-вариант.
# Суффикс допускаем И кириллицей, И латиницей («000012а» / «000012a»): у клиентов
# с латинскими картами «Перечень» тоже набран латиницей. Раньше класс был только
# кириллическим — из-за этого латинские «а»-строки не проходили фильтр и ЦЕЛИКОМ
# выпадали из 6_4 вместе со своими людьми (Sud-tibbiyot: строка «000012a» на
# 5 чел./2 жен. терялась, Morfalogiya давала 7/6 вместо 12/8).
_PERECHEN_WP_RE = re.compile(r"^\d{4,6}[а-яёА-ЯЁa-zA-Z]?$")


def _perechen_header_columns(grid: list[list[str]]) -> dict[str, int]:
    """Индексы колонок Перечня по объединённому заголовку (строки 0–3)."""
    ncols = max(len(r) for r in grid) if grid else 0
    header = [" ".join(grid[r][c] if c < len(grid[r]) else "" for r in range(min(4, len(grid))))
              for c in range(ncols)]
    cols: dict[str, int] = {}
    for i, h in enumerate(header):
        hl = h.lower()
        if "code" not in cols and M.PERECHEN_COL_CODE in hl:
            cols["code"] = i
        if "workplace" not in cols and M.PERECHEN_COL_WORKPLACE in hl:
            cols["workplace"] = i
        if "position" not in cols and M.PERECHEN_COL_POSITION in hl and M.PERECHEN_COL_CODE not in hl:
            cols["position"] = i
        if "workers" not in cols and M.PERECHEN_COL_WORKERS in hl:
            cols["workers"] = i
        if "female" not in cols and M.PERECHEN_COL_FEMALE in hl and M.PERECHEN_COL_WORKERS not in hl:
            cols["female"] = i
    cols.setdefault("workplace", 0)
    cols.setdefault("position", 1)
    cols.setdefault("code", 2)
    cols.setdefault("workers", 3)
    cols.setdefault("female", 4)
    return cols


def parse_perechen(docx_path: str | Path) -> dict[str, dict]:
    """Вернуть {workplace_no: {'job_code', 'position'}} из «Перечня»."""
    doc = read_docx(docx_path)
    grid = max(doc.tables, key=len) if doc.tables else None
    if not grid:
        return {}

    cols = _perechen_header_columns(grid)
    workplace_col, position_col, code_col = cols["workplace"], cols["position"], cols["code"]

    result: dict[str, dict] = {}
    for row in grid:
        if workplace_col >= len(row):
            continue
        wp = normalize_spaces(row[workplace_col])
        if not _PERECHEN_WP_RE.match(wp):
            continue
        result[wp] = {
            "job_code": normalize_spaces(row[code_col]) if code_col < len(row) else "",
            "position": normalize_spaces(row[position_col]) if position_col < len(row) else "",
        }
    return result


# --- 6_4: позиции Перечня с привязкой к подразделению (строки-разделители) --
def parse_perechen_positions_6_4(docx_path: str | Path) -> tuple[list[dict], list[str]]:
    """Вернуть позиции Перечня по порядку документа для сборки 6_4.

    Каждая позиция: ``{workplace_no, subdivision, employees_count,
    female_count}``. Подразделение определяется строками-разделителями —
    НАСТОЯЩИМИ горизонтальными слияниями ячеек на всю ширину таблицы (все
    ячейки строки указывают на один и тот же ``<w:tc>``), а не эвристикой по
    тексту — это надёжно отличает их от обычных строк данных независимо от
    формулировки названия подразделения (см. Шаг 2 спецификации). Повторы
    одного номера (напр. две карты одной "а"-позиции на разные смены)
    сохраняются как отдельные позиции, не схлопываются в одну (Шаг 3/5).
    """
    document = Document(str(docx_path))
    warnings: list[str] = []
    table = max(document.tables, key=lambda t: len(t.rows), default=None)
    if table is None:
        return [], ["«Перечень»: таблица позиций не найдена (6_4)."]

    grid = [[normalize_spaces(c.text) for c in row.cells] for row in table.rows]
    cols = _perechen_header_columns(grid)
    workplace_col, workers_col, female_col = cols["workplace"], cols["workers"], cols["female"]

    positions: list[dict] = []
    current_sub = ""
    for row in table.rows:
        cells = row.cells
        text0 = normalize_spaces(cells[0].text) if cells else ""
        is_merged_full_row = len({id(c._tc) for c in cells}) == 1
        if is_merged_full_row and text0 and not _PERECHEN_WP_RE.match(text0):
            current_sub = text0
            continue
        wp = normalize_spaces(cells[workplace_col].text) if workplace_col < len(cells) else ""
        if not _PERECHEN_WP_RE.match(wp):
            continue
        positions.append({
            "workplace_no": wp,
            "subdivision": current_sub,
            "employees_count": normalize_number(cells[workers_col].text) if workers_col < len(cells) else "",
            "female_count": normalize_number(cells[female_col].text) if female_col < len(cells) else "",
        })
    return positions, warnings
