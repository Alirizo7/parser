"""Заполнение шаблонов 5_1б, 6_4, 6_5 и 6_6 через python-docx.

Подход: берём фиксированный пустой шаблон-ассет, очищаем строки тела ниже
шапки и генерируем тело заново из единого датасета (см. ``pipeline``).
Форматирование сохраняем, клонируя «строку-прототип» тела шаблона.
"""
from __future__ import annotations

import re
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Pt
from docx.table import Table

from openpyxl import load_workbook

from . import mapping as M
from . import xlsx
from .extract import injury_risk_value, split_workplace_no, workplace_sort_key
from .normalize import (
    _to_int,
    class_rank,
    fold,
    fold_contains,
    max_class,
    normalize_spaces,
    to_cyrillic,
    to_latin,
)

ASSETS_DIR = Path(__file__).resolve().parent.parent / "assets"
TEMPLATE_5_1B = ASSETS_DIR / "template_5_1b.docx"
TEMPLATE_6_5 = ASSETS_DIR / "template_6_5.docx"
TEMPLATE_6_4 = ASSETS_DIR / "template_6_4.docx"
TEMPLATE_6_6 = {
    "cyr": ASSETS_DIR / "template_6_6_cyr.docx",
    "lat": ASSETS_DIR / "template_6_6_lat.docx",
}
TEMPLATE_EXCEL = {n: ASSETS_DIR / f"template_excel_{n}.xlsx" for n in range(1, 6)}


def _transliterate_doc(doc, lang: str) -> None:
    """Привести ВЕСЬ текст документа к выбранному письму перед сохранением.

    ``lang == 'lat'`` → транслитерируем кириллицу в латиницу (заголовки шаблона
    и данные); ``'cyr'`` → ничего не делаем (текущее поведение, документ уже на
    кириллице). Цифры/коды/пунктуацию ``to_latin`` не трогает.
    """
    if lang != "lat":
        return

    def fix_paragraph(p):
        for run in p.runs:
            if run.text:
                run.text = to_latin(run.text)

    def fix_table(tbl):
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    fix_paragraph(p)
                for inner in cell.tables:
                    fix_table(inner)

    for p in doc.paragraphs:
        fix_paragraph(p)
    for tbl in doc.tables:
        fix_table(tbl)


# --- Низкоуровневые помощники работы с ячейками -----------------------------
def set_cell_text(cell, text: str) -> None:
    """Записать текст в ячейку, сохранив форматирование первого run-а."""
    text = "" if text is None else str(text)
    p = cell.paragraphs[0]
    if p.runs:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(text)


def _clear_body(table: Table, keep_header_rows: int):
    """Удалить строки тела ниже шапки, вернув копию строки-прототипа."""
    rows = list(table.rows)
    prototype = deepcopy(rows[keep_header_rows]._tr) if len(rows) > keep_header_rows else None
    for row in rows[keep_header_rows:]:
        table._tbl.remove(row._tr)
    return prototype


def _append_row(table: Table, prototype):
    """Добавить новую строку тела по прототипу; вернуть объект строки."""
    tr = deepcopy(prototype)
    table._tbl.append(tr)
    return table.rows[-1]


def _prevent_row_split(row) -> None:
    """Не позволять Word/LibreOffice разрывать строку таблицы между страницами."""
    if row._tr.xpath("./w:trPr/w:cantSplit"):
        return
    row._tr.get_or_add_trPr().append(OxmlElement("w:cantSplit"))


def _repeat_table_header(row) -> None:
    """Повторять строку шапки таблицы на каждой следующей странице."""
    if row._tr.xpath("./w:trPr/w:tblHeader"):
        return
    row._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))


def _keep_row_with_next(row) -> None:
    """Не отрывать строку-заголовок/промежуточный итог от следующей строки."""
    seen: set[int] = set()
    for cell in row.cells:
        if id(cell._tc) in seen:
            continue
        seen.add(id(cell._tc))
        for paragraph in cell.paragraphs:
            p_pr = paragraph._p.get_or_add_pPr()
            if not p_pr.xpath("./w:keepNext"):
                p_pr.append(OxmlElement("w:keepNext"))


def _remove_trailing_empty_paragraphs(doc) -> None:
    """Удалить пустые абзацы ассета после последней таблицы.

    В шаблоне 6.4 после тела осталось несколько пустых абзацев. Когда таблица
    заканчивается у нижнего поля страницы, LibreOffice переносит их на отдельную
    полностью пустую страницу. Содержательных абзацев не удаляем.
    """
    body = doc._element.body
    children = list(body.iterchildren())
    last_table = max(
        (index for index, child in enumerate(children) if child.tag.endswith("}tbl")),
        default=-1,
    )
    for child in children[last_table + 1:]:
        if not child.tag.endswith("}p"):
            continue
        text = "".join(child.itertext()).strip()
        if not text:
            body.remove(child)


def _remove_content_after_table(doc, table: Table) -> None:
    """Удалить служебный хвост шаблона после указанной таблицы.

    Ассет 6_5 был получен из готового клиентского файла и содержал после
    сводной таблицы скан подписей конкретной организации. Одного удаления
    абзаца недостаточно: картинка останется внутри DOCX как доступное вложение.
    Поэтому вместе с XML-узлом удаляем его image relationship из пакета.
    """
    body = doc._element.body
    children = list(body.iterchildren())
    try:
        table_index = children.index(table._tbl)
    except ValueError:
        return

    for child in children[table_index + 1:]:
        if child.tag.endswith("}sectPr"):
            continue
        image_rel_ids = child.xpath(".//a:blip/@r:embed")
        body.remove(child)
        for rel_id in image_rel_ids:
            if rel_id in doc.part.rels:
                doc.part.drop_rel(rel_id)


def _pct_num(value: str) -> str:
    m = re.search(r"\d+", value or "")
    return m.group(0) if m else ""


# --- 5_1б: вредные вещества по рабочим местам -------------------------------
def group_substances(substances: list[dict]) -> list[tuple[str, str]]:
    """Сгруппировать вещества по проценту воздействия.

    Возвращает список (имена_через_запятую, процент) в порядке появления
    процентов. Вещества внутри группы уже упорядочены на этапе извлечения.
    """
    groups: dict[str, list[str]] = {}
    for s in substances:
        pct = _pct_num(s.get("pct", ""))
        groups.setdefault(pct, []).append(s["name"])

    # Как в эталоне: первое вещество в ячейке с заглавной, остальные — строчными
    # («Углерод оксиди, азот оксиди, силикат чанги (лой)»).
    def join(names: list[str]) -> str:
        if not names:
            return ""
        return ", ".join([names[0]] + [n[:1].lower() + n[1:] for n in names[1:]])

    return [(join(names), pct) for pct, names in groups.items()]


def render_5_1b(workplaces: list[dict], out_path: str | Path,
                *, template_path: str | Path = TEMPLATE_5_1B, lang: str = "cyr") -> Path:
    """Сформировать документ 5_1б из датасета рабочих мест."""
    doc = Document(str(template_path))
    table = doc.tables[0]
    # Шапка: R0 (заголовки) + R1 («1|2|3»). Тело — с R2.
    prototype = _clear_body(table, keep_header_rows=2)

    # Гарантируем порядок: 000011 < 000011а < 000012 (даже если вход не отсортирован)
    workplaces = sorted(workplaces, key=lambda w: workplace_sort_key(w.get("workplace_no", "")))
    for wp in workplaces:
        groups = group_substances(wp.get("substances", []))
        if not groups:
            # Веществ в карте нет (раздел 1.1 пуст) — выводим РАБОЧЕЕ МЕСТО ВСЁ
            # РАВНО, оставляя «Модданинг номи» пустым (флаг substances_missing).
            # Так число строк совпадает с ожиданием, а пропуск виден в документе.
            row = _append_row(table, prototype)
            set_cell_text(row.cells[0], wp["workplace_no"])
            set_cell_text(row.cells[1], "")
            set_cell_text(row.cells[2], "")
            continue
        for gi, (names, pct) in enumerate(groups):
            row = _append_row(table, prototype)
            set_cell_text(row.cells[0], wp["workplace_no"] if gi == 0 else "")
            set_cell_text(row.cells[1], names)
            set_cell_text(row.cells[2], f"{pct} %" if pct else "")

    _transliterate_doc(doc, lang)
    out_path = Path(out_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return out_path


# --- 6_5: большая сводная таблица (25 логических колонок = 26 grid) ----------
# Соответствие grid-колонок таблицы 1 полям записи рабочего места.
# «Сут» (молоко) занимает 2 объединённые grid-колонки c22–c23.
_FACTOR_COLS = [
    (3, "chem"), (4, "biological"), (5, "aerosols"), (6, "noise"),
    (7, "infrasound"), (8, "ultrasound_air"), (9, "vibration_general"),
    (10, "vibration_local"), (11, "em_field"), (12, "ionizing"),
    (13, "microclimate"), (14, "lighting"), (15, "severity"),
    (16, "intensity"), (17, "overall"),
]

_REQS_LABELS = {
    "name": "корхонанинг номи",
    "parent": "юқори турувчи",
    "address": "юридик манзили",
    "product": "асосий маҳсулот",
}
_REQS_CODES = (("stir", "стир"), ("ifut", "ифут"), ("mxbt", "мхбт"), ("mxbt", "mxбt"))


def row_values_6_5(rec: dict) -> list[str]:
    """26 значений grid-строки сводной таблицы для одной записи."""
    f = rec.get("factors", {}) or {}
    b = rec.get("benefits", {}) or {}
    vals = [""] * 26
    vals[0] = rec.get("workplace_no", "")
    # Должность в 6_5 — из «Перечня» (как в эталонах клиента); иначе из карты
    vals[1] = (
        rec.get("position_from_perechen_6_5")
        or rec.get("position_from_perechen")
        or rec.get("position", "")
    )
    vals[2] = rec.get("job_code_6_5") or rec.get("job_code", "")
    for ci, key in _FACTOR_COLS:
        vals[ci] = f.get(key, "-") or "-"
    # Травмоопасность — из п.2.3 карты (то же поле, что считает 6_4); эвристика
    # «медик→1, иначе→2» осталась лишь фолбэком. См. extract.injury_risk_value.
    vals[18] = injury_risk_value(rec)
    vals[19] = rec.get("ppe_provided", "")
    vals[20] = b.get("extra_leave", "")
    vals[21] = b.get("reduced_hours", "")
    vals[22] = b.get("milk", "")
    vals[23] = b.get("milk", "")  # «Сут» объединена на 2 grid-колонки
    vals[24] = b.get("therapeutic_food", "")
    vals[25] = rec.get("privileged_pension", "")
    return vals


def _norm(s: str) -> str:
    return " ".join((s or "").split())


_CODE_CELL_KEYS = {"стир": "stir", "ифут": "ifut", "мхбт": "mxbt", "mxбt": "mxbt"}


def _fill_reqs(table, company: dict) -> None:
    """Подставить реквизиты компании в таблицу 0.

    ВСЕГДА перезаписываем все ячейки значений из источника; если значение не
    извлеклось — ОЧИЩАЕМ ячейку (пусто), чтобы данные компании-примера из
    шаблона (СТИР/ИФУТ/МХБТ Бухоро) не «протекали» к другому клиенту.
    """
    for row in table.rows:
        cells = row.cells
        c0 = cells[0].text
        # Текстовые поля «метка | значение»: значение в объединённой ячейке cells[2]
        for key, anchor in _REQS_LABELS.items():
            if fold_contains(c0, anchor) and len(cells) > 2:
                set_cell_text(cells[2], company.get(key, "") or "")
        # Строка кодов: ячейка-метка СТИР/ИФУТ/МХБТ → следующая ячейка
        for i, cell in enumerate(cells):
            fc = fold(cell.text)
            for cyr, key in _CODE_CELL_KEYS.items():
                if fc == fold(cyr) and i + 1 < len(cells):
                    set_cell_text(cells[i + 1], company.get(key, "") or "")


def _group_by_subdivision(workplaces: list[dict]) -> list[tuple[str, list[dict]]]:
    """Сгруппировать соседние рабочие места по подразделению (как в эталоне)."""
    groups: list[tuple[str, list[dict]]] = []
    for wp in workplaces:
        sub = wp.get("subdivision", "") or "—"
        if not groups or groups[-1][0] != sub:
            groups.append((sub, []))
        groups[-1][1].append(wp)
    return groups


def _group_by_subdivision_6_5(workplaces: list[dict]) -> list[tuple[str, list[dict]]]:
    """Сгруппировать 6_5 по физическим блокам Перечня.

    Название подразделения из карты остаётся фолбэком для архивов без
    Перечня. Когда Перечень разобран, его заголовок является авторитетным и не
    дробит один блок из-за регистра или опечатки в отдельных картах.
    """
    groups: list[tuple[str, list[dict]]] = []
    for wp in workplaces:
        sub = wp.get("subdivision_6_5") or wp.get("subdivision") or "—"
        if not groups or groups[-1][0] != sub:
            groups.append((sub, []))
        groups[-1][1].append(wp)
    return groups


def _format_data_row_6_5(row) -> None:
    """Применить формат тела 6_5 из заполненных клиентских эталонов."""
    seen: set[int] = set()
    position_tc = id(row.cells[1]._tc) if len(row.cells) > 1 else None
    for cell in row.cells:
        marker = id(cell._tc)
        if marker in seen:
            continue
        seen.add(marker)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for paragraph in cell.paragraphs:
            paragraph.alignment = (
                WD_ALIGN_PARAGRAPH.LEFT if marker == position_tc else WD_ALIGN_PARAGRAPH.CENTER
            )
            for run in paragraph.runs:
                run.bold = False
                run.font.size = Pt(9)


def render_6_5(company_data: dict, workplaces: list[dict], out_path: str | Path,
               *, template_path: str | Path = TEMPLATE_6_5, lang: str = "cyr") -> Path:
    """Сформировать сводный документ 6_5 из датасета.

    Прототипы строк:
    * R3 — строка-заголовок группы (одна объединённая на всю ширину ячейка);
    * R2 — строка нумерации «1|2|…|25»: у неё объединение «Сут» стоит на ВЕРНОЙ
      позиции (grid-колонки 22–23). У пустых строк данных шаблона (R4) объединение
      смещено на 21–22, из-за чего «сокр.день» и «Сут» схлопывались в одну ячейку
      (баг проявлялся, когда их значения различались). Поэтому строки данных
      клонируем из R2 и перезаписываем числа-плейсхолдеры реальными значениями.
    """
    doc = Document(str(template_path))
    _fill_reqs(doc.tables[0], company_data)
    summary = doc.tables[1]

    # При наличии Перечня повторяем его физический порядок. Карты, которых в
    # Перечне нет, не теряем и ставим после него в обычном порядке номеров.
    def order_6_5(wp: dict):
        order = wp.get("perechen_order_6_5")
        if isinstance(order, int):
            return (0, order, workplace_sort_key(wp.get("workplace_no", "")))
        return (1, 10**9, workplace_sort_key(wp.get("workplace_no", "")))

    workplaces = sorted(workplaces, key=order_6_5)

    rows = list(summary.rows)
    group_proto = deepcopy(rows[3]._tr)  # заголовок группы (спанящая ячейка)
    data_proto = deepcopy(rows[2]._tr)   # строка данных (объединение «Сут» на 22–23)
    for row in rows[3:]:                 # очищаем тело ниже 3-уровневой шапки
        summary._tbl.remove(row._tr)

    for sub, members in _group_by_subdivision_6_5(workplaces):
        header = _append_row(summary, group_proto)
        set_cell_text(header.cells[0], sub)
        _prevent_row_split(header)
        _keep_row_with_next(header)
        for wp in members:
            row = _append_row(summary, data_proto)
            cells = row.cells
            for ci, val in enumerate(row_values_6_5(wp)):
                if ci < len(cells):
                    set_cell_text(cells[ci], val)
            _format_data_row_6_5(row)
            _prevent_row_split(row)

    # Не переносим в новый документ скан подписей организации-примера,
    # оставшийся в исходном шаблоне после сводной таблицы.
    _remove_content_after_table(doc, summary)
    _transliterate_doc(doc, lang)
    out_path = Path(out_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return out_path


# --- 6_6: план мероприятий по вредным факторам -----------------------------
# Порядок строк повторяет порядок факторов в готовых клиентских 6.6. Одна
# группа даёт максимум одно мероприятие на карту, даже если вредны несколько
# родственных подфакторов (например, шум и вибрация).
_PLAN_GROUPS_6_6 = (
    ("chem", ("chem", "aerosols")),
    ("biological", ("biological",)),
    ("physical", ("noise", "infrasound", "ultrasound_air", "vibration_general", "vibration_local")),
    ("radiation", ("em_field", "ionizing")),
    ("microclimate", ("microclimate",)),
    ("lighting", ("lighting",)),
    ("severity", ("severity",)),
    ("intensity", ("intensity",)),
)

_PLAN_MEASURES_6_6 = {
    "lat": {
        "chem": (
            "Zararli mehnat sharoitlarini hisobga olgan holda, shaxsiy nafas olish "
            "vositalaridan foydalaning. Yuqori samaradorlikni saqlash va mehnat "
            "unumdorligini oshirish uchun tartibga solinadigan tanaffuslarni ta‘minlang. "
            "Ish va dam olish tartibini saqlang."
        ),
        "biological": (
            "Biologik omillar ta’siri davomiyligini kamaytiring, sanitariya-gigiyena "
            "tadbirlarini kuchaytiring va xodimlarni zarur himoya vositalari bilan ta’minlang."
        ),
        "physical": (
            "Shovqinning va tebranishning zararli ta‘sirini kamaytirish uchun hodimga "
            "mos keladi. Eshitish organlari va tebranishni kamaytirish choralari "
            "(eshitish vositasi yoki qabul qiluvchi va maxsus tebranishni oldini oluvchi "
            "maxsus oyoq kiyimlar yoki rezina poliklar)."
        ),
        "radiation": (
            "Nurlanish ta’sirini kamaytiring: himoya ekranlari va xavfsiz masofani "
            "ta’minlang, ta’sir vaqtini me’yorlang, davriy nazorat o‘tkazing hamda "
            "xodimlarni belgilangan shaxsiy himoya vositalari bilan ta’minlang."
        ),
        "microclimate": (
            "Sanoat va turar-joy binolarida normal issiqlik sharoitlari va mikroiqlim, "
            "toza havoni ta‘minlash uchun issiqlik va havo pardalari, aspiratsiya va "
            "chang va gazni ushlab turuvchi qurilmalarni yangi isitish va shamollatish "
            "tizimlarini o‘rnatish va mavjudlarini rekonstruksiya qilish."
        ),
        "lighting": (
            "Yoritishni me’yorga keltiring (chiroqlar sonini ko‘paytiring, ish joyining "
            "ustiga umumiy yoritish chiroqini o‘rnating, mahalliy yoritish chiroqini "
            "o‘rnating, yuqori quvvatli lampalardan foydalaning)"
        ),
        "intensity": (
            "Mehnat jarayonining asabiy-emotsional keskinligini hisobga olgan holda, "
            "xodimning ichki mehnat rejimini tartibga solish va mehnat jarayonining "
            "hissiy tanglik darajasini pasaytirish bo‘yicha chora-tadbirlarni ishlab "
            "chiqish va qabul qilish."
        ),
    },
    "cyr": {
        "chem": (
            "Зарарли меҳнат шароитларини ҳисобга олган ҳолда, шахсий нафас олиш "
            "воситаларидан фойдаланинг. Юқори самарадорликни сақлаш ва меҳнат "
            "унумдорлигини ошириш учун тартибга солинадиган танаффусларни таъминланг. "
            "Иш ва дам олиш тартибини сақланг"
        ),
        "biological": (
            "Биологик омиллар таъсири давомийлигини камайтиринг, санитария-гигиена "
            "тадбирларини кучайтиринг ва ходимларни зарур ҳимоя воситалари билан таъминланг."
        ),
        "physical": (
            "Шовқиннинг зарарли таъсирини камайтириш учун ходимга мос келади "
            "Эшитиш органлари (ешитиш воситаси ёки қабул қилувчи)"
        ),
        "radiation": (
            "Нурланиш таъсирини камайтиринг: ҳимоя экранлари ва хавфсиз масофани "
            "таъминланг, таъсир вақтини меъёрланг, даврий назорат ўтказинг ҳамда "
            "ходимларни белгиланган шахсий ҳимоя воситалари билан таъминланг."
        ),
        "microclimate": (
            "Саноат ва турар-жой биноларида нормал иссиқлик шароитлари ва микроиқлим, "
            "тоза ҳавони таъминлаш учун иссиқлик ва ҳаво пардалари, аспирация ва чанг "
            "ва газни ушлаб турувчи қурилмаларни янги иситиш ва шамоллатиш тизимларини "
            "ўрнатиш ва мавжудларини реконструкция қилиш."
        ),
        "lighting": (
            "Ёритишни меъёрга келтиринг (чироқлар сонини кўпайтиринг, иш жойининг "
            "устига умумий ёритиш чироқини ўрнатинг, маҳаллий ёритиш чироқини "
            "ўрнатинг, юқори қувватли лампалардан фойдаланинг)"
        ),
        "intensity": (
            "Меҳнат жараёнининг асабий-эмоционал кескинлигини ҳисобга олган ҳолда, "
            "ходимнинг ички меҳнат режимини тартибга солиш ва меҳнат жараёнининг "
            "ҳиссий танглик даражасини пасайтириш бўйича чора-тадбирларни ишлаб "
            "чиқиш ва қабул қилиш."
        ),
    },
}

_SEVERITY_MEASURE_6_6 = {
    "lat": (
        "Mehnatning og‘irligini hisobga olgan holda{detail} ish va dam olish rejimini "
        "ishlab chiqish tavsiya etiladi, gimnastika mashqlari bilan ish kuni davomida "
        "uzoq muddatli tartibga solingan tanaffuslarni nazarda tutadi. Ish kuni va "
        "haftaning dinamikasida mehnat jarayonining og‘irligini kamaytirish uchun "
        "mehnat va dam olishning oqilona almashinuvi rejimiga qat‘iy rioya qilish kerak."
    ),
    "cyr": (
        "Меҳнатнинг оғирлигини ҳисобга олган ҳолда{detail} иш ва дам олиш режимини "
        "ишлаб чиқиш тавсия этилади, гимнастика машқлари билан иш куни давомида "
        "узоқ муддатли тартибга солинган танаффусларни назарда тутади. Иш куни ва "
        "ҳафтанинг динамикасида меҳнат жараёнининг оғирлигини камайтириш учун "
        "меҳнат ва дам олишнинг оқилона алмашинуви режимига қатъий риоя қилиш керак."
    ),
}

_PLAN_EXECUTION_6_6 = {
    "lat": ("Qo'llanma", "Tashkilot byudjeti", "MM va TB xodimiga",
            "Doimiy ravishda", "Boshqaruv"),
    # В двух кириллических эталонах колонка «масъул» тоже равна «Қўлланма».
    "cyr": ("Қўлланма", "Ташкилот бюджети", "Қўлланма",
            "Доимий равишда", "Бошқарув"),
}


def _plan_text_6_6(text: str, lang: str) -> str:
    """Локализовать только управляемый статичный текст 6_6.

    Данные компании и должности не транслитерируем: имена/бренды/коды должны
    оставаться ровно такими, какими были извлечены или исправлены оператором.
    """
    return to_cyrillic(text) if lang == "cyr" else text


def _group_is_harmful_6_6(workplace: dict, group: str, keys: tuple[str, ...], lang: str) -> bool:
    factors = workplace.get("factors", {}) or {}
    if group == "chem" and lang == "cyr":
        # Оба кириллических эталона не создают мероприятие, когда 3.x получен
        # только расчётной строкой «Йиғиш омили»; нужен конкретный вредный агент
        # либо вредный аэрозоль. Латинский эталон, напротив, включает этот случай.
        substances = workplace.get("substances", []) or []
        harmful_agent = any(
            class_rank(item.get("cls", "")) >= 30
            and "yigish omili" not in fold(item.get("name", ""))
            for item in substances
        )
        return harmful_agent or class_rank(factors.get("aerosols", "")) >= 30
    return any(class_rank(factors.get(key, "")) >= 30 for key in keys)


def _severity_measure_6_6(workplace: dict, lang: str) -> str:
    posture = ((workplace.get("plan_6_6") or {}).get("severity_posture") or "").strip()
    if posture:
        posture = to_cyrillic(posture) if lang == "cyr" else to_latin(posture)
    detail = f" ({posture})" if posture else ""
    return _SEVERITY_MEASURE_6_6[lang].format(detail=detail)


def _workplace_plan_rows_6_6(workplace: dict, lang: str) -> list[list[str]]:
    lang = lang if lang in TEMPLATE_6_6 else "cyr"
    operational = _PLAN_EXECUTION_6_6[lang]
    result: list[list[str]] = []
    first = True
    for group, keys in _PLAN_GROUPS_6_6:
        if not _group_is_harmful_6_6(workplace, group, keys, lang):
            continue
        measure = (
            _severity_measure_6_6(workplace, lang)
            if group == "severity"
            else _PLAN_MEASURES_6_6[lang][group]
        )
        card = _plan_text_6_6(f"Karta\n№{workplace.get('workplace_no', '')}", lang) if first else ""
        first = False
        result.append([
            card, measure,
            operational[0], operational[1], operational[2], operational[3], operational[4], "",
        ])
    return result


def plan_sections_6_6(workplaces: list[dict], *, lang: str = "cyr") -> list[tuple[str, list[list[str]]]]:
    """Сгруппировать строки плана по подразделениям, как в готовых 6.6."""
    sections: list[tuple[str, list[list[str]]]] = []
    ordered = sorted(workplaces, key=lambda w: workplace_sort_key(w.get("workplace_no", "")))
    for workplace in ordered:
        rows = _workplace_plan_rows_6_6(workplace, lang)
        if not rows:
            continue
        subdivision = workplace.get("subdivision", "") or "—"
        if not sections or sections[-1][0] != subdivision:
            sections.append((subdivision, []))
        sections[-1][1].extend(rows)
    return sections


def plan_rows_6_6(workplaces: list[dict], *, lang: str = "cyr") -> list[list[str]]:
    """Плоский список карточных строк 6.6 (без строк подразделений)."""
    rows = [row for _subdivision, block in plan_sections_6_6(workplaces, lang=lang) for row in block]
    if not rows:
        rows.append([
            "—", _plan_text_6_6("3–4-sinf zararli omillar aniqlanmadi.", lang),
            "", "", "", "", "", "",
        ])
    return rows


def _append_spanning_row_6_6(table: Table, prototype, text: str):
    row = _append_row(table, prototype)
    _prevent_row_split(row)
    cell = row.cells[0].merge(row.cells[-1])
    # ``merge`` переносит по абзацу из каждой из восьми ячеек. Обычный
    # ``set_cell_text`` очистил бы runs, но оставил семь пустых абзацев и
    # искусственно раздувал строку. Присваивание ``cell.text`` схлопывает их.
    cell.text = text
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.bold = True
    return row


def _format_data_row_6_6(row, lang: str) -> None:
    """Привести строку к фактической вёрстке трёх готовых форм.

    В исходных пустых шаблонах у строки-прототипа остался
    межабзацный интервал 14 pt, из-за чего каждое мероприятие
    занимало лишнее место. Эталоны используют 9 pt для кириллицы
    и 10 pt для латиницы, без дополнительных интервалов.
    """
    font_size = Pt(9 if lang == "cyr" else 10)
    for cell_index, cell in enumerate(row.cells):
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            for run in paragraph.runs:
                run.font.size = font_size
                if cell_index == 0:
                    run.bold = True


def _ppe_summary_6_6(company_data: dict, lang: str) -> str:
    mandatory = (company_data.get("ppe_mandatory_6_6") or "").strip()
    additional = (company_data.get("ppe_additional_6_6") or "").strip()
    # В эталонах дополнительный коллективный договор переносится только в
    # кириллической форме и только когда в карте указан настоящий реквизит
    # (дата/номер), а не пустая подпись «Жамоа шартномаси».
    if lang != "cyr" or not re.search(r"\d", additional):
        additional = ""
    values = [mandatory, additional]
    return " ".join(v for v in values if v)


def render_6_6(company_data: dict, workplaces: list[dict], out_path: str | Path,
               *, template_path: str | Path | None = None, lang: str = "cyr") -> Path:
    """Сформировать приложение 6.6 — динамический план мероприятий."""
    lang = lang if lang in TEMPLATE_6_6 else "cyr"
    selected_template = Path(template_path) if template_path else TEMPLATE_6_6[lang]
    doc = Document(str(selected_template))
    # Первая таблица — чистый двухколоночный блок согласования/утверждения;
    # реквизиты и план всегда две последние таблицы шаблона.
    _fill_reqs(doc.tables[-2], company_data)
    measures = doc.tables[-1]
    for header_row in measures.rows[:2]:
        _repeat_table_header(header_row)
    prototype = _clear_body(measures, keep_header_rows=2)
    if prototype is None:
        raise ValueError("Шаблон 6_6 не содержит строку-прототип тела")

    sections = plan_sections_6_6(workplaces, lang=lang)
    if sections:
        for subdivision, rows in sections:
            _append_spanning_row_6_6(measures, prototype, subdivision)
            for values in rows:
                row = _append_row(measures, prototype)
                for cell, value in zip(row.cells, values):
                    set_cell_text(cell, value)
                _format_data_row_6_6(row, lang)
    else:
        for values in plan_rows_6_6(workplaces, lang=lang):
            row = _append_row(measures, prototype)
            for cell, value in zip(row.cells, values):
                set_cell_text(cell, value)
            _format_data_row_6_6(row, lang)

    ppe_summary = _ppe_summary_6_6(company_data, lang)
    if ppe_summary:
        _append_spanning_row_6_6(measures, prototype, ppe_summary)

    out_path = Path(out_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return out_path


# --- 6_4: сводная қайднома (итоги по подразделениям) ------------------------
# Таблица 1 шаблона (13 grid-колонок):
#   c0  — название строки/подразделения
#   c1  — итого (аттестовано рабочих мест / занято людей / из них женщин)
#   c2..c5  — по классу условий труда: 1, 2, 3, 4
#   c6..c8  — по классу травмоопасности: 1, 2, 3
#   c9..c11 — обеспеченность ЯТҲВ: Мос / Мос эмас / Ятҳв кўзда тутилмаган
#   c12 — попадает под «3-4 класс и/или ЯТҲВ не соответствует»
#
# Каждый блок (итог по компании, итог по подразделению) — 4 строки: строка-
# заголовок (название) + 3 строки данных (иш ўринлари / ходимлар / аёллар).
_ROW_KINDS = ("units", "employees", "female")

# Позиция считается «учтённой в 6_4», только если пайплайн сопоставил её со
# строкой Перечня (см. pipeline._pair_perechen_positions_6_4) — иначе у неё
# нет ключа subdivision_6_4, и включать её в подсчёт нельзя (см. mapping/extract).


def _overall_degree(wp: dict) -> str:
    """Старшая цифра общего класса («3.2» → «3»); "" если класс не извлечён."""
    overall = (wp.get("factors") or {}).get("overall", "") or ""
    return overall[0] if overall[:1] in ("1", "2", "3", "4") else ""


def _aggregate_group_6_4(positions: list[dict], warnings: list[str]) -> dict[str, list[int]]:
    """3 строки (units/employees/female) × 13 grid-значений (c0 не считаем).

    Источник числа работников/женщин — Перечень (``employees_count_6_4``/
    ``female_count_6_4``, см. Шаг 4 спецификации), НЕ карта: карта даёт лишь
    3 итоговых показателя (класс условий труда, класс травмоопасности,
    статус ЯТҲВ) для распределения этих чисел по колонкам.
    """
    out: dict[str, list[int]] = {}
    for row_kind in _ROW_KINDS:
        vals = [0] * 13
        for wp in positions:
            if row_kind == "units":
                # «а»-суффиксные строки Перечня — доп. смена/условие ТОГО ЖЕ
                # рабочего места, не новая единица (сверено с эталоном: сумма
                # строк БЕЗ суффикса == «Иш ўринлари, бирлик» итога компании).
                _, suffix = split_workplace_no(wp.get("workplace_no", ""))
                w = 0 if suffix else 1
            elif row_kind == "employees":
                w = _to_int(wp.get("employees_count_6_4"))
            else:
                w = _to_int(wp.get("female_count_6_4"))
            if not w:
                continue
            degree = _overall_degree(wp)
            risk = wp.get("injury_risk_class_6_4", "")
            ppe = wp.get("ppe_status_6_4", "")

            vals[1] += w
            if degree == "1":
                vals[2] += w
            elif degree == "2":
                vals[3] += w
            elif degree == "3":
                vals[4] += w
            elif degree == "4":
                vals[5] += w
            elif row_kind == "units":
                warnings.append(
                    f"{wp.get('workplace_no', '?')}: класс условий труда не извлечён из карты — "
                    "не учтён в разбивке по классам 6_4."
                )
            if risk == "1":
                vals[6] += w
            elif risk == "2":
                vals[7] += w
            elif risk == "3":
                vals[8] += w
            elif row_kind == "units":
                warnings.append(
                    f"{wp.get('workplace_no', '?')}: класс травмоопасности не извлечён из карты — "
                    "не учтён в разбивке по классам 6_4."
                )
            # c9/c10 — соответствие требованиям ЯТҲВ (п.3.3); c11 — «СИЗ не
            # предусмотрены» (таблица СИЗ п.3.2). Это НЕЗАВИСИМЫЕ величины:
            # рабочее место одновременно «Мос» (c9) и «кўзда тутилмаган» (c11).
            # В эталоне Мос=114 (все) и кўзда=93 (подмножество) — суммы намеренно
            # пересекаются, поэтому НЕ раскладываем по взаимоисключающим колонкам.
            if ppe == "mos_emas":
                vals[10] += w
            else:
                vals[9] += w
            if wp.get("ppe_not_envisaged_6_4"):
                vals[11] += w
            # c12 («3-4 даража ва/ёки Ятҳв мос эмас») — в эталоне ПУСТОЙ по всем
            # подразделениям (автор оставил его незаполненным), поэтому c12 не
            # заполняем (vals[12] остаётся 0 → рендерится «-»).
        out[row_kind] = vals
    return out


def _group_positions_by_subdivision_6_4(
    positions: list[dict],
) -> list[tuple[list[str], list[dict]]]:
    """Собрать блоки 6_4 строго по физическим сериям строк «Перечня».

    Новый датасет несёт ``subdivision_group_6_4`` и полный список иерархических
    заголовков. Для сохранённых старых датасетов используем соседние серии, но
    никогда не объединяем одинаковые названия глобально через весь документ.
    """
    groups: list[tuple[list[str], list[dict]]] = []
    keys: list[tuple] = []
    for wp in positions:
        headers = list(wp.get("subdivision_headers_6_4") or [])
        if not headers:
            headers = [wp.get("subdivision_6_4", "") or "—"]
        explicit_group = wp.get("subdivision_group_6_4")
        key = ("id", explicit_group) if explicit_group is not None else (
            "legacy", tuple(fold(h) for h in headers)
        )
        if not groups or keys[-1] != key:
            groups.append((headers, []))
            keys.append(key)
        groups[-1][1].append(wp)
    return groups


def _fill_numeric_row(row, vals: list[int]) -> None:
    """Заполнить строку данных (c1..c12); c0 — метка, не трогаем.

    Ноль отображаем как «-» (эталон использует прочерк для отсутствующих
    значений в каждой колонке-категории, не «0»).
    """
    cells = row.cells
    for i in range(1, min(13, len(cells))):
        set_cell_text(cells[i], "-" if not vals[i] else str(vals[i]))
    seen: set[int] = set()
    for cell in cells:
        if id(cell._tc) in seen:
            continue
        seen.add(id(cell._tc))
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _format_group_heading_6_4(row, heading: str) -> None:
    """Формат заголовка блока как в готовых 6.4.

    Верхнеуровневые названия, набранные прописными, — жирные прямые; обычные
    подразделения/участки — жирный курсив. Оба варианта центрированы.
    """
    letters = [char for char in heading if char.isalpha()]
    is_upper_level = bool(letters) and all(char.isupper() for char in letters)
    cell = row.cells[0]
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.bold = True
            run.italic = not is_upper_level


def _normalize_acronym_6_4(doc, lang: str) -> None:
    """Привести профессиональную аббревиатуру к эталонному регистру."""
    replacement = "YaTHV" if lang == "lat" else "ЯТҲВ"
    pattern = re.compile(r"yathv" if lang == "lat" else r"ятҳв", re.IGNORECASE)
    for table in doc.tables:
        for row in table.rows:
            seen: set[int] = set()
            for cell in row.cells:
                if id(cell._tc) in seen:
                    continue
                seen.add(id(cell._tc))
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.text = pattern.sub(replacement, run.text)


def render_6_4(company_data: dict, workplaces: list[dict], out_path: str | Path,
               *, template_path: str | Path = TEMPLATE_6_4, lang: str = "cyr",
               warnings: list[str] | None = None) -> Path:
    """Сформировать сводную қайднома 6_4 (итоги по подразделениям) из датасета.

    Статическую шапку и прототипы строк берём из ассета, но тело всегда строим
    заново по потоку «Перечня». Это сохраняет родительские/дочерние заголовки,
    повторные одноимённые физические блоки и исходный порядок без зависимости
    от названий подразделений компании, использованной в шаблоне.
    """
    if warnings is None:
        warnings = []
    doc = Document(str(template_path))
    _fill_reqs(doc.tables[0], company_data)
    # Во всех изученных готовых 6.4 значения реквизитов жирные и центрированы.
    reqs = doc.tables[0]
    for row in reqs.rows[:4]:
        cell = row.cells[2]
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                if run.font.size is None:
                    run.font.size = Pt(11)
    if len(reqs.rows) > 4:
        for index in (1, 3, 5):
            if index >= len(reqs.rows[4].cells):
                continue
            cell = reqs.rows[4].cells[index]
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.bold = True
    summary = doc.tables[1]

    positions = [wp for wp in workplaces if "subdivision_6_4" in wp]
    positions.sort(key=lambda w: (
        w.get("perechen_order_6_4", 10**9),
        workplace_sort_key(w.get("workplace_no", "")),
    ))

    rows = list(summary.rows)
    total_data_rows = rows[4:7]  # «Корхона бўйича жами»: 3 строки данных

    # Итоговый блок по компании — заполняем на месте (заголовок R3 не трогаем)
    totals = _aggregate_group_6_4(positions, warnings)
    for row_obj, kind in zip(total_data_rows, _ROW_KINDS):
        _fill_numeric_row(row_obj, totals[kind])

    groups = _group_positions_by_subdivision_6_4(positions)
    header_proto = deepcopy(rows[10]._tr) if len(rows) > 10 else None
    data_protos = [deepcopy(r._tr) for r in rows[11:14]]

    # Готовые клиентские документы разделяют общий итог (7 строк) и перечень
    # подразделений (повторная 3-строчная шапка) на отдельные таблицы.
    subdivision_tbl = deepcopy(summary._tbl)
    for tr in list(subdivision_tbl.tr_lst)[:7]:
        subdivision_tbl.remove(tr)
    # В готовых 6.4 общий итог всегда завершает первую страницу, а повторная
    # шапка подразделений начинает следующую. Явный разрыв не оставляет шапку
    # одиноко внизу первой страницы без первого блока данных.
    page_break = OxmlElement("w:p")
    p_pr = OxmlElement("w:pPr")
    p_pr.append(OxmlElement("w:pageBreakBefore"))
    spacing = OxmlElement("w:spacing")
    spacing.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}before", "0")
    spacing.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}after", "0")
    spacing.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}line", "1")
    spacing.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}lineRule", "exact")
    p_pr.append(spacing)
    page_break.append(p_pr)
    summary._tbl.addnext(page_break)
    page_break.addnext(subdivision_tbl)
    subdivisions = Table(subdivision_tbl, summary._parent)
    for row_obj in list(summary.rows)[7:]:
        summary._tbl.remove(row_obj._tr)
    for row_obj in list(subdivisions.rows)[3:]:
        subdivisions._tbl.remove(row_obj._tr)

    for header_row in subdivisions.rows[:3]:
        _prevent_row_split(header_row)

    if header_proto is not None and len(data_protos) == 3:
        for headers, members in groups:
            for heading in headers or ["—"]:
                row_obj = _append_row(subdivisions, header_proto)
                set_cell_text(row_obj.cells[0], heading)
                _format_group_heading_6_4(row_obj, heading)
                _prevent_row_split(row_obj)
                _keep_row_with_next(row_obj)
            agg = _aggregate_group_6_4(members, warnings)
            for index, (proto, kind) in enumerate(zip(data_protos, _ROW_KINDS)):
                row_obj = _append_row(subdivisions, proto)
                _fill_numeric_row(row_obj, agg[kind])
                _prevent_row_split(row_obj)
                if index < 2:
                    _keep_row_with_next(row_obj)

    _remove_trailing_empty_paragraphs(doc)

    _transliterate_doc(doc, lang)
    _normalize_acronym_6_4(doc, lang)
    out_path = Path(out_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return out_path


# ===========================================================================
# Excel-протоколы лабораторных замеров (5 файлов) — render_excel_1..5
# ===========================================================================
# Тело каждого протокола генерируется из единого датасета клонированием прото-
# блока рабочего места (см. services/xlsx.py). Управляемая часть ассета держится
# на письме эталона (латиница); в конце приводится к output_lang «к целевому
# письму» (xlsx.transliterate_region). Статичная шапка (приборы/НД/лаборатория) —
# «как есть», меняется только заказчик (company_data).

# Геометрия ассетов (совпадает с build_excel_assets.CFG). ncols — число колонок.
EXCEL_CFG = {
    1: dict(ncols=12, managed_start=27, group_row=30, block_start=31, block_len=1),
    2: dict(ncols=11, managed_start=27, group_row=30, block_start=31, block_len=4),
    3: dict(ncols=11, managed_start=24, group_row=27, block_start=28, block_len=9),
    4: dict(ncols=11, managed_start=24, group_row=27, block_start=28, block_len=5),
    5: dict(ncols=11, managed_start=24, group_row=27, block_start=28, block_len=16),
}

# Метки строки-заказчика в шапке (двуязычные: файл 1 — узб. латиница, 2–5 — рус.)
_EXCEL_REQS_NAME = ("buyurtmachi nomi", "наименование заказчик")
_EXCEL_REQS_ADDR = ("buyurtmachi manzil", "адрес заказчик")


def _fill_excel_reqs(ws, company: dict) -> None:
    """Вписать имя/адрес заказчика в шапку (перезаписать метку-строку).

    Значение идёт ИНЛАЙН после метки в той же ячейке («Наименование заказчика:
    <имя>»); длинное имя в эталоне переносится на следующую строку — её очищаем,
    чтобы имя компании-образца не «протекло» (анти-утечка, как в docx _fill_reqs).
    """
    company = company or {}
    name_row = None
    for r in range(1, EXCEL_HEADER_SCAN + 1):
        cell = ws.cell(row=r, column=1)
        v = cell.value
        if not isinstance(v, str):
            continue
        if any(fold_contains(v, a) for a in _EXCEL_REQS_NAME):
            cell.value = _reqs_prefix(v) + (company.get("name", "") or "")
            name_row = r
        elif any(fold_contains(v, a) for a in _EXCEL_REQS_ADDR):
            cell.value = _reqs_prefix(v) + (company.get("address", "") or "")
    # Очистить строку-продолжение длинного имени образца (анти-утечка), но НЕ
    # трогать её, если это уже другая метка (адрес и т.п. — имя было в одну строку).
    if name_row is not None:
        cont = ws.cell(row=name_row + 1, column=1)
        cv = cont.value
        is_label = isinstance(cv, str) and any(
            fold_contains(cv, a) for a in (*_EXCEL_REQS_NAME, *_EXCEL_REQS_ADDR)
        )
        if not is_label:
            cont.value = None


EXCEL_HEADER_SCAN = 26  # в пределах скольких строк искать метки заказчика


def _reqs_prefix(text: str) -> str:
    """Префикс метки до двоеточия включительно («Наименование заказчика: »)."""
    return (text.split(":", 1)[0] + ": ") if ":" in text else (text + " ")


# --- Преобразование значений замеров в ячейки -------------------------------
def _d_norma(value):
    """Норма в колонку D: число / диапазон-строка / «-»; пусто → None."""
    if value is None:
        return None
    s = str(value).strip()
    if s == "":
        return None
    if s == "-":
        return "-"
    return xlsx.to_number(s)


def _cls_cell(value):
    """Класс в ячейку: число (2.0→2, 3.1→3.1); «-»/пусто → None."""
    if value is None:
        return None
    s = str(value).strip()
    if s in ("", "-"):
        return None
    return xlsx.to_number(s)


# Колонки D..J (норма, 1/2/3-фаолият, факт, время, класс) для файлов 2–5
_DJ_COLS = range(4, 11)


def _measure_fill(m: dict | None, *, with_class: bool = True) -> dict:
    """Заполнение колонок D..J под-строки по замеру {norma,actual,time,cls}.

    * замера нет → пустая рамка (D..J очищены);
    * факт-текст «йўқ» → псевдозамеры E/F/G и факт H = тот же текст, время/класс пусты;
    * числовой факт → норма в D, факт в H, формулы E/F/G из прото (KEEP), время/класс.

    ``with_class=False`` — не выводить класс (J): в некоторых под-строках эталон
    его не сообщает (напр. естественная освещённость КЕО в файле 4).
    """
    if not m:
        return {c: xlsx.CLEAR for c in _DJ_COLS}
    actual = xlsx.to_number(m.get("actual"))
    if actual is None:
        return {c: xlsx.CLEAR for c in _DJ_COLS}
    d = _d_norma(m.get("norma"))
    if isinstance(actual, str):  # факт — текст «йўқ»
        return {4: d, 5: actual, 6: actual, 7: actual, 8: actual, 9: xlsx.CLEAR, 10: xlsx.CLEAR}
    cls = _cls_cell(m.get("cls")) if with_class else xlsx.CLEAR
    return {4: d, 5: xlsx.KEEP, 6: xlsx.KEEP, 7: xlsx.KEEP, 8: actual,
            9: xlsx.to_number(m.get("time")), 10: cls}


def _label_d_fill(text: str) -> dict:
    """Под-строка «метка в D, остальное пусто» (Ishlar toifasi / Vizual ish / энергия)."""
    fill = {c: xlsx.CLEAR for c in _DJ_COLS}
    if text:
        fill[4] = text
    return fill


# --- Построители под-строк блока (файлы 2–5) --------------------------------
def _subrows_file2(wp: dict) -> list[dict]:
    pm = wp.get("physical_measurements") or {}
    return [
        _measure_fill(pm.get("noise")),
        _measure_fill(pm.get("vibration_local")),
        _measure_fill(pm.get("vibration_general")),
        _measure_fill(pm.get("infrasound")),
    ]


def _subrows_file3(wp: dict) -> list[dict]:
    mc = wp.get("microclimate_measurements") or {}
    return [
        _label_d_fill(""),                       # 0 энергозатраты (только K=итог)
        _label_d_fill(mc.get("category_label") or ""),  # 1 Ishlar toifasi
        _measure_fill(mc.get("temp")),           # 2 температура
        _measure_fill(mc.get("air_speed")),      # 3 скорость воздуха
        _measure_fill(mc.get("humidity")),       # 4 влажность
        _measure_fill(mc.get("heat_radiation")), # 5 теплоизлучение («йўқ»/число)
        _measure_fill(mc.get("outdoor_temp")),   # 6 температура (откр. территория), 1.10.1
        _label_d_fill(""),                       # 7 WBGT
        _label_d_fill(""),                       # 8 средняя тепл. нагрузка TNS
    ]


def _subrows_file4(wp: dict) -> list[dict]:
    lg = wp.get("lighting_measurements") or {}
    return [
        _label_d_fill(lg.get("discharge") or ""),          # 0 разряд зрит. работ
        _measure_fill(lg.get("natural"), with_class=False),  # 1 естественная (КЕО): без класса
        _measure_fill(lg.get("combined")),                 # 2 смешанная (КЕО)
        _measure_fill(lg.get("artificial")),               # 3 искусственная (лк)
        # 4 пульсация: эталон систематически НЕ сообщает её в этом протоколе
        # (0/52 РМ), хотя в картах раздел 1.11.6 заполнен — оставляем пустой рамкой,
        # чтобы вывод совпал с эталоном по данным.
        _measure_fill(None),
    ]


# Файл 5: 16 под-строк; данные конкретной 1.4.x лежат в верхней строке пары.
# ⚠️ Прото-формулы псевдозамеров (E/F/G = H±0.01) есть в эталоне ТОЛЬКО на строке
# 1.4.10 (единственной заполненной у медиков). Если у клиента заполнен другой
# раздел (1.4.1–1.4.9), его норма/факт/время/класс выводятся корректно, но
# декоративные колонки «1/2/3-фаолият» останутся пустыми (в прото-строке формулы
# нет). Данные при этом не теряются; при появлении таких клиентов формулы можно
# до-синтезировать. То же касается пульсации в файле 4.
_EM_SUBROW_SECTION = {
    0: "1.4.1", 1: "1.4.2", 2: "1.4.3", 3: "1.4.4", 4: "1.4.5",
    6: "1.4.6", 8: "1.4.7", 10: "1.4.8", 12: "1.4.9", 14: "1.4.10",
}


def _subrows_file5(wp: dict) -> list[dict]:
    em = wp.get("em_measurements") or {}
    fills = []
    for i in range(16):
        sec = _EM_SUBROW_SECTION.get(i)
        fills.append(_measure_fill(em.get(sec)) if sec else _label_d_fill(""))
    return fills


# --- Итоговый класс (K/L) по файлу ------------------------------------------
def _factor(wp: dict, key: str) -> str:
    return (wp.get("factors") or {}).get(key, "-") or "-"


def _final_file1(wp):
    return _cls_cell(_factor(wp, "chem"))


def _final_file2(wp):
    phys = [_factor(wp, k) for k in
            ("noise", "infrasound", "ultrasound_air", "vibration_general", "vibration_local")]
    return _cls_cell(max_class(phys))


def _final_file3(wp):
    return _cls_cell(_factor(wp, "microclimate"))


def _final_file4(wp):
    return _cls_cell(_factor(wp, "lighting"))


def _final_file5(wp):
    return _cls_cell(_factor(wp, "em_field"))


def _has_em(wp: dict) -> bool:
    return any(v for v in (wp.get("em_measurements") or {}).values())


# В исходном тёплом шаблоне строка наружной температуры была пустой, а строка
# теплоизлучения содержала литералы «йўқ». Для числовых замеров им нужны формулы,
# аналогичные соседним строкам. Добавляем их в прототип в памяти ДО capture_row;
# текстовый факт затем штатно переопределит E:G/H через _measure_fill.
_MICROCLIMATE_PROTO_FORMULAS = {
    5: ("+1", "+2", "-3"),          # теплоизлучение, факт H
    6: ("+0.1", "+0.2", "-0.3"),  # температура открытой территории, факт H
}


def _seed_microclimate_proto_formulas(ws, block_start: int) -> None:
    for offset, suffixes in _MICROCLIMATE_PROTO_FORMULAS.items():
        row = block_start + offset
        for column, suffix in zip(range(5, 8), suffixes):
            ws.cell(row=row, column=column).value = f"=H{row}{suffix}"


# --- Общий рендер файлов 2–5 (фикс. блок из N под-строк) --------------------
def _render_excel_blocks(company_data, workplaces, out_path, *, idx, subrows_of,
                         final_of, include, grouped, lang):
    cfg = EXCEL_CFG[idx]
    ncols, gr, bs, blen = cfg["ncols"], cfg["group_row"], cfg["block_start"], cfg["block_len"]
    wb = load_workbook(str(TEMPLATE_EXCEL[idx]))
    ws = wb["complete"]
    if idx == 3:
        _seed_microclimate_proto_formulas(ws, bs)

    # Снять прототипы ДО очистки тела
    group_proto = xlsx.capture_row(ws, gr, ncols) if gr else None
    group_merges = xlsx.capture_merges(ws, gr, gr) if gr else []
    block_proto = [xlsx.capture_row(ws, bs + i, ncols) for i in range(blen)]
    block_merges = xlsx.capture_merges(ws, bs, bs + blen - 1)

    body_start = gr if gr else bs
    xlsx.clear_body(ws, body_start)
    _fill_excel_reqs(ws, company_data)

    wps = [w for w in sorted(workplaces, key=lambda w: workplace_sort_key(w.get("workplace_no", "")))
           if include(w)]

    cur = body_start

    def emit_block(wp):
        nonlocal cur
        fills = subrows_of(wp)
        start = cur
        for i in range(blen):
            f = dict(fills[i])
            if i == 0:  # № РМ, должность, итоговый класс — только 1-я строка (merge)
                f[1] = wp.get("workplace_no", "")
                f[2] = wp.get("position_from_perechen") or wp.get("position", "")
                f[ncols] = final_of(wp)
            xlsx.emit_row(ws, cur, block_proto[i], ncols, f)
            cur += 1
        xlsx.apply_merges(ws, block_merges, start)  # A/B/K спаны + внутренние объединения

    if grouped:
        for sub, members in _group_by_subdivision(wps):
            xlsx.emit_row(ws, cur, group_proto, ncols, {1: sub})
            grow = cur
            cur += 1
            xlsx.apply_merges(ws, group_merges, grow)
            for wp in members:
                emit_block(wp)
    else:
        for wp in wps:
            emit_block(wp)

    xlsx.transliterate_region(ws, lang, cfg["managed_start"])
    out_path = Path(out_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    wb.save(str(out_path))
    return out_path


def render_excel_1(company_data, workplaces, out_path, *,
                   template_path=None, lang: str = "cyr") -> Path:
    """Файл 1 — вредные вещества (одна строка на вещество, блок переменной высоты)."""
    cfg = EXCEL_CFG[1]
    ncols, gr, bs = cfg["ncols"], cfg["group_row"], cfg["block_start"]
    wb = load_workbook(str(TEMPLATE_EXCEL[1]))
    ws = wb["complete"]

    group_proto = xlsx.capture_row(ws, gr, ncols)
    group_merges = xlsx.capture_merges(ws, gr, gr)      # A30:L30
    sub_proto = xlsx.capture_row(ws, bs, ncols)         # одна прото-строка вещества

    xlsx.clear_body(ws, gr)
    _fill_excel_reqs(ws, company_data)

    wps = sorted(workplaces, key=lambda w: workplace_sort_key(w.get("workplace_no", "")))
    cur = gr

    for sub, members in _group_by_subdivision(wps):
        xlsx.emit_row(ws, cur, group_proto, ncols, {1: sub})
        grow = cur
        cur += 1
        xlsx.apply_merges(ws, group_merges, grow)
        for wp in members:
            subs = wp.get("substances") or []
            position = wp.get("position_from_perechen") or wp.get("position", "")
            final = _final_file1(wp)
            start = cur
            if not subs:  # РМ без веществ — одна пустая строка (не теряем место)
                xlsx.emit_row(ws, cur, sub_proto, ncols, {
                    1: wp.get("workplace_no", ""), 2: position, 3: xlsx.CLEAR,
                    4: xlsx.CLEAR, 5: xlsx.CLEAR, 6: xlsx.CLEAR, 7: xlsx.CLEAR,
                    8: xlsx.CLEAR, 9: xlsx.CLEAR, 10: xlsx.CLEAR, 11: xlsx.CLEAR, 12: final,
                })
                cur += 1
            for si, s in enumerate(subs):
                actual = xlsx.to_number(s.get("actual"))  # I — факт (число или текст)
                if isinstance(actual, str):
                    # Факт — текст («йўқ» и т.п.): НЕ оставляем формулы «=I±шаг» на
                    # текстовой ячейке (иначе Excel даст #VALUE!) — как _measure_fill
                    # для файлов 2–5: псевдозамеры F/G/H = тот же текст.
                    pseudo = {6: actual, 7: actual, 8: actual}
                else:
                    pseudo = {6: xlsx.KEEP, 7: xlsx.KEEP, 8: xlsx.KEEP}  # формулы =I±шаг
                f = {
                    3: s.get("name", ""),           # C — вещество
                    4: xlsx.CLEAR,                  # D — класс опасности (нет в карте)
                    5: _d_norma(s.get("norma")),    # E — ПДК/норма
                    **pseudo,                       # F/G/H
                    9: actual,                      # I — факт
                    10: xlsx.to_number(s.get("pct")),  # J — время
                    11: _cls_cell(s.get("cls")),       # K — класс вещества
                }
                if si == 0:
                    f[1] = wp.get("workplace_no", "")
                    f[2] = position
                    f[12] = final              # L — итог по РМ
                else:
                    f[1] = xlsx.CLEAR
                    f[2] = xlsx.CLEAR
                    f[12] = xlsx.CLEAR
                xlsx.emit_row(ws, cur, sub_proto, ncols, f)
                cur += 1
            # Вертикальные объединения № РМ / должность / итог по высоте блока
            xlsx.merge_span(ws, start, cur - 1, 1)
            xlsx.merge_span(ws, start, cur - 1, 2)
            xlsx.merge_span(ws, start, cur - 1, 12)

    xlsx.transliterate_region(ws, lang, cfg["managed_start"])
    out_path = Path(out_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    wb.save(str(out_path))
    return out_path


def render_excel_2(company_data, workplaces, out_path, *, template_path=None, lang="cyr") -> Path:
    """Файл 2 — физические факторы (шум/вибрация/инфразвук)."""
    return _render_excel_blocks(company_data, workplaces, out_path, idx=2,
                                subrows_of=_subrows_file2, final_of=_final_file2,
                                include=lambda w: True, grouped=True, lang=lang)


def render_excel_3(company_data, workplaces, out_path, *, template_path=None, lang="cyr") -> Path:
    """Файл 3 — микроклимат."""
    return _render_excel_blocks(company_data, workplaces, out_path, idx=3,
                                subrows_of=_subrows_file3, final_of=_final_file3,
                                include=lambda w: True, grouped=True, lang=lang)


def render_excel_4(company_data, workplaces, out_path, *, template_path=None, lang="cyr") -> Path:
    """Файл 4 — освещённость."""
    return _render_excel_blocks(company_data, workplaces, out_path, idx=4,
                                subrows_of=_subrows_file4, final_of=_final_file4,
                                include=lambda w: True, grouped=True, lang=lang)


def render_excel_5(company_data, workplaces, out_path, *, template_path=None, lang="cyr") -> Path:
    """Файл 5 — магнитные поля / ЭМИ (по отделам; только РМ с замером 1.4.x)."""
    return _render_excel_blocks(company_data, workplaces, out_path, idx=5,
                                subrows_of=_subrows_file5, final_of=_final_file5,
                                include=_has_em, grouped=True, lang=lang)
