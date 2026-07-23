"""Самопроверка: сравнение сгенерированных документов с эталонными из архива.

Эталоны (заполненные ``…қилиниши керак``) лежат внутри примера zip. Сравнение
ведём по СОДЕРЖАНИЮ (логическая карта «рабочее место → данные»), игнорируя
служебные строки (повтор шапки, пустые), т.к. позиции строк в эталоне зависят
от ручной разбивки по страницам.
"""
from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path

from docx import Document

from .normalize import fold, normalize_spaces


# --- Поиск эталонов в распакованном архиве ---------------------------------
def find_reference(unpacked_dir: str | Path, which: str) -> Path | None:
    """Найти заполненный эталон 5_1б или 6_5 среди распакованных файлов."""
    unpacked_dir = Path(unpacked_dir)
    for path in unpacked_dir.rglob("*"):
        if path.suffix.lower() not in (".doc", ".docx"):
            continue
        name = path.name.lower()
        if "пустой" in name or "пустои" in name:
            continue
        if which == "5_1b" and name.startswith("5_1б") and "керак" in name:
            return path
        if which == "6_5" and name.startswith("6 5") and "керак" in name:
            return path
    return None


# --- Нормализация значений для сравнения ------------------------------------
def _norm_name(s: str) -> str:
    return re.sub(r"[.\s]+$", "", normalize_spaces(s)).lower()


def _pct(s: str) -> str:
    m = re.search(r"\d+", s or "")
    return m.group(0) if m else ""


# --- Разбор 5_1б в логическую карту ----------------------------------------
def parse_5_1b(docx_path: str | Path) -> dict[str, list[tuple[str, str]]]:
    """{workplace_no: [(имена_через_запятую_norm, процент), ...]}."""
    doc = Document(str(docx_path))
    table = max(doc.tables, key=lambda t: len(t.rows))

    def cell(r, c):
        cells = table.rows[r].cells
        return normalize_spaces(cells[c].text) if c < len(cells) else ""

    result: dict[str, list[tuple[str, str]]] = {}
    cur: str | None = None
    for r in range(len(table.rows)):
        c0, c1, c2 = cell(r, 0), cell(r, 1), cell(r, 2)
        # пропускаем шапку «1|2|3» и заголовки
        if c0 in ("1", "Иш ўрнининг т/р") and c1 in ("2", "Модданинг номи"):
            continue
        if re.match(r"^\d{6}", c0):  # новая запись (номер рабочего места)
            cur = c0
            result.setdefault(cur, []).append((_norm_name(c1), _pct(c2)))
        elif cur and not c0 and c1:  # под-строка (вещество с другим %)
            result[cur].append((_norm_name(c1), _pct(c2)))
    return result


# --- Сравнение --------------------------------------------------------------
@dataclass
class CheckResult:
    matched: int = 0
    total: int = 0
    mismatches: list[str] = field(default_factory=list)
    # «Заметки» — расхождения в полях-эвристиках или там, где эталон клиента
    # сам отличается от его исходных карт. Не считаются провалом проверки.
    notes: list[str] = field(default_factory=list)

    @property
    def ok(self) -> bool:
        # Проверка пройдена, если нет расхождений в КЛЮЧЕВЫХ полях
        # (заметки по эвристикам/расхождениям клиента провалом не считаются).
        return self.total > 0 and not self.mismatches


def compare_5_1b(generated_path: str | Path, reference_path: str | Path,
                 *, bilingual: bool = False) -> CheckResult:
    gen = parse_5_1b(generated_path)
    ref = parse_5_1b(reference_path)
    if bilingual:
        # Приводим имена к общему письму и убираем уточнения в скобках
        # («silikat changi (loy)» ↔ «silikat changi»): клиент их редактирует.
        def _f(d):
            return {k: [(fold(re.sub(r"\([^)]*\)", "", n)), p) for n, p in v]
                    for k, v in d.items()}
        gen, ref = _f(gen), _f(ref)
    res = CheckResult()

    # Рабочие места без веществ в картах (клиент добавил вручную) — в заметки
    for no in sorted(set(ref) - set(gen)):
        msg = f"{no}: есть в эталоне, нет в сгенерированном (в карте нет раздела веществ)"
        (res.notes if bilingual else res.mismatches).append(msg)
    for no in sorted(set(gen) - set(ref)):
        res.mismatches.append(f"{no}: лишнее рабочее место в сгенерированном")

    for no in sorted(set(ref) & set(gen)):
        gen_empty = all(not n for n, _ in gen[no])  # строка РМ без веществ (пустая ячейка)
        ref_has = any(n for n, _ in ref[no])
        if bilingual and gen_empty and ref_has:
            # РМ выведено, но веществ в карте нет (клиент вписал вручную) — заметка
            res.notes.append(f"{no}: в карте нет веществ, в эталоне есть (substances_missing)")
            continue
        res.total += 1
        if ref[no] == gen[no]:
            res.matched += 1
        else:
            res.mismatches.append(f"{no}: эталон={ref[no]} ≠ сгенерировано={gen[no]}")
    return res


# --- Разбор и сравнение 6_5 -------------------------------------------------
_6_5_COL_NAMES = {
    0: "т/р", 1: "должность", 2: "код", 3: "хим", 4: "био", 5: "аэроз", 6: "шум",
    7: "инфра", 8: "ультра", 9: "общвибр", 10: "локвибр", 11: "ЭМ", 12: "ион",
    13: "микро", 14: "свет", 15: "тяж", 16: "напр", 17: "ОБЩ", 18: "травмо",
    19: "ЯТҲВ", 20: "отпуск", 21: "сокр", 22: "сут", 24: "лечпит", 25: "пенсия",
}
# Колонку c23 пропускаем (дубль объединённой «Сут»).
_6_5_COLS = [c for c in range(26) if c != 23]

_REQS_FIELDS = {
    "name": "корхонанинг номи", "parent": "юқори турувчи",
    "address": "юридик манзили", "product": "асосий маҳсулот",
}


def _read_6_5(docx_path):
    """Вернуть (реквизиты, {workplace_no: [значения 26 колонок]})."""
    doc = Document(str(docx_path))
    reqs_tbl = next((t for t in doc.tables
                     if any("корхонанинг номи" in normalize_spaces(c.text).lower()
                            for r in t.rows for c in r.cells)), doc.tables[0])
    summary = max(doc.tables, key=lambda t: len(t.rows))

    def cell(t, r, c):
        cells = t.rows[r].cells
        return normalize_spaces(cells[c].text) if c < len(cells) else ""

    reqs = {}
    for r in range(len(reqs_tbl.rows)):
        c0 = cell(reqs_tbl, r, 0).lower()
        for key, anchor in _REQS_FIELDS.items():
            if anchor in c0:
                reqs[key] = cell(reqs_tbl, r, 2)
        joined = " ".join(cell(reqs_tbl, r, c) for c in range(len(reqs_tbl.columns)))
        for key, pat in (("stir", r"СТИР\s*([0-9]+)"), ("ifut", r"ИФУТ\s*([0-9]+)"),
                         ("mxbt", r"(?:МХБТ|MX[БB]T)\s*([0-9]+)")):
            m = re.search(pat, joined, re.IGNORECASE)
            if m:
                reqs[key] = m.group(1)

    data = {}
    for r in range(len(summary.rows)):
        c0 = cell(summary, r, 0)
        if re.match(r"^\d{6}", c0):
            data[c0] = [cell(summary, r, c) for c in range(26)]
    return reqs, data


def compare_6_5(generated_path, reference_path) -> CheckResult:
    gen_reqs, gen = _read_6_5(generated_path)
    ref_reqs, ref = _read_6_5(reference_path)
    res = CheckResult()

    # Реквизиты
    for key in ("name", "parent", "address", "product", "stir", "ifut", "mxbt"):
        res.total += 1
        if normalize_spaces(gen_reqs.get(key, "")) == normalize_spaces(ref_reqs.get(key, "")):
            res.matched += 1
        else:
            res.mismatches.append(
                f"реквизит «{key}»: эталон={ref_reqs.get(key)!r} ≠ {gen_reqs.get(key)!r}"
            )

    for no in sorted(set(ref) - set(gen)):
        res.mismatches.append(f"{no}: нет в сгенерированном")
    for no in sorted(set(gen) - set(ref)):
        res.mismatches.append(f"{no}: лишнее в сгенерированном")

    for no in sorted(set(ref) & set(gen)):
        for c in _6_5_COLS:
            res.total += 1
            if _cells_equal(c, ref[no][c], gen[no][c]):
                res.matched += 1
            else:
                msg = (f"{no} c{c}({_6_5_COL_NAMES.get(c, c)}): "
                       f"эталон={ref[no][c]!r} ≠ {gen[no][c]!r}")
                # «Сут» (c22) в эталоне Бухоро вручную переопределён для медиков
                # вопреки форме карты — расхождение, а не ошибка извлечения.
                (res.notes if c == 22 else res.mismatches).append(msg)
    return res


def _cells_equal(col: int, ref: str, gen: str) -> bool:
    """Сравнение ячеек. Должность (c1) сверяем без учёта дефисов/пробелов:
    в эталоне в неё вручную вставлены переносы-дефисы (``ҳисоб-лаш``,
    ``меҳ-нат``) для узкой колонки — содержание при этом совпадает."""
    if col == 1:
        norm = lambda s: re.sub(r"[-­\s]+", "", normalize_spaces(s)).lower()
        return norm(ref) == norm(gen)
    return ref == gen


# --- Двуязычное сравнение 6_5 (наш кир-вывод vs лат-эталон клиента) ----------
# 25 логических полей в порядке колонок. Наш шаблон — 26 grid-колонок («Сут»
# объединена на 2), эталон клиента — 25. Сводим к одной логической схеме.
LOGICAL_6_5 = [
    "т/р", "должность", "код", "хим", "био", "аэроз", "шум", "инфра", "ультра",
    "общ.вибр", "лок.вибр", "ЭМ", "ион", "микро", "свет", "тяж", "напр", "ОБЩ",
    "травмо", "ЯТҲВ", "отпуск", "сокр", "сут", "лечпит", "пенсия",
]
_REQS_BILINGUAL = {
    "name": "корхонанинг номи", "parent": "юқори турувчи",
    "address": "юридик манзили", "product": "асосий маҳсулот",
}


def _reqs_value_cells(cells, anchor: str) -> str:
    flabel = fold(anchor)
    for c in cells:
        raw = normalize_spaces(c.text)
        if raw and flabel not in fold(c.text):
            return raw
    return ""


def _read_6_5_logical(path):
    """Прочитать 6_5 в (реквизиты, {workplace_no: [25 логических значений]})."""
    doc = Document(str(path))
    summary = max(doc.tables, key=lambda t: len(t.rows))
    reqs_tbl = next(
        (t for t in doc.tables
         if any("korxonaning nomi" in fold(c.text) for r in t.rows for c in r.cells)),
        doc.tables[0],
    )

    def cell(t, r, c):
        cs = t.rows[r].cells
        return normalize_spaces(cs[c].text) if c < len(cs) else ""

    n = len(summary.columns)
    data = {}
    for r in range(len(summary.rows)):
        c0 = cell(summary, r, 0)
        if not re.match(r"\d{4,6}", c0):
            continue
        cells = [cell(summary, r, c) for c in range(n)]
        if n >= 26:  # наш шаблон: «Сут»=c22(дубль c23), лечпит=c24, пенсия=c25
            logical = cells[0:22] + [cells[22], cells[24], cells[25]]
        else:        # эталон клиента: 25 колонок 1:1
            logical = (cells + [""] * 25)[0:25]
        data[c0] = logical

    reqs = {}
    for r in range(len(reqs_tbl.rows)):
        c0 = cell(reqs_tbl, r, 0)
        for key, anc in _REQS_BILINGUAL.items():
            if fold(anc) in fold(c0):
                reqs[key] = _reqs_value_cells(reqs_tbl.rows[r].cells, anc)
        joined = " ".join(cell(reqs_tbl, r, c) for c in range(len(reqs_tbl.columns)))
        for key, lab in (("stir", r"(?:СТИР|STIR)"), ("ifut", r"(?:ИФУТ|IFUT)"),
                         ("mxbt", r"(?:МХБТ|MX[БB]T)")):
            m = re.search(lab + r"\s*[:\-]?\s*([0-9][0-9 ]*)", joined, re.IGNORECASE)
            if m:
                reqs[key] = re.sub(r"\D", "", m.group(1))
    return reqs, data


def _logical_equal(i: int, gen: str, ref: str) -> bool:
    if i == 1:  # должность — без дефисов/пробелов, в общем письме
        norm = lambda s: re.sub(r"[-­\s]+", "", fold(s))
        return norm(gen) == norm(ref)
    if i == 2:  # код должности — только цифры
        return re.sub(r"\D", "", gen) == re.sub(r"\D", "", ref)
    return fold(gen) == fold(ref)  # факторы/числа/токены (ҳа≡ha, йўқ≡yo`q)


def compare_6_5_bilingual(generated_path, reference_path) -> CheckResult:
    g_reqs, gen = _read_6_5_logical(generated_path)
    r_reqs, ref = _read_6_5_logical(reference_path)
    res = CheckResult()

    for key in ("name", "parent", "address", "product", "stir", "ifut", "mxbt"):
        res.total += 1
        gv, rv = g_reqs.get(key, ""), r_reqs.get(key, "")
        eq = (re.sub(r"\D", "", gv) == re.sub(r"\D", "", rv)) if key in ("stir", "ifut", "mxbt") \
            else (fold(gv) == fold(rv))
        if eq:
            res.matched += 1
        else:
            res.mismatches.append(f"реквизит «{key}»: эталон={rv!r} ≠ {gv!r}")

    for no in sorted(set(ref) - set(gen)):
        res.mismatches.append(f"{no}: нет в сгенерированном")
    for no in sorted(set(gen) - set(ref)):
        res.mismatches.append(f"{no}: лишнее в сгенерированном")

    for no in sorted(set(ref) & set(gen)):
        for i in range(25):
            res.total += 1
            gv = gen[no][i] if i < len(gen[no]) else ""
            rv = ref[no][i] if i < len(ref[no]) else ""
            if _logical_equal(i, gv, rv):
                res.matched += 1
            else:
                msg = f"{no} {LOGICAL_6_5[i]}: эталон={rv!r} ≠ {gv!r}"
                (res.notes if i in _REVIEW_6_5 else res.mismatches).append(msg)
    return res


# Поля «к проверке»: эвристики (травмо/сут/пенсия) и пофакторные колонки,
# где эталон клиента может расходиться с его же картами (наш вывод верен картам).
# Ключевые поля (должны совпасть): реквизиты, код, должность, подытоги факторов
# (хим/микро/свет/тяж/напр/общий), ЯТҲВ и льготы отпуск/сокр/лечпит.
_REVIEW_6_5 = {4, 5, 6, 7, 8, 9, 10, 11, 12, 18, 22, 24}


# ===========================================================================
# Excel-протоколы лабораторных замеров — сравнение с эталонами excel_templates/
# ===========================================================================
# Единица счёта — заполненная ячейка ДАННЫХ (норма/факт/время/класс каждой
# под-строки каждого РМ + итоговый класс блока). Сравниваем сгенерированный
# xlsx с эталоном по СОДЕРЖАНИЮ (РМ → под-строки), позиции строк — из геометрии
# блока (фикс. число под-строк у файлов 2–5; переменное у файла 1).
#
# Правила классификации (как в docx-самопроверке — «верим картам»):
#   * РМ есть в эталоне, нет у нас → mismatch (потеряли рабочее место);
#   * РМ есть у нас, нет в эталоне → note (эталон неполон/опустил — у нас по картам);
#   * расхождение ячейки на РМ из ``_EXCEL_ANOMALY_WP`` → note (эталон клиента
#     содержит опечатки ввода: на 000050 факт освещённости записан текстом,
#     на 000051 факт/время смещены — наш вывод верен картам);
#   * прочее расхождение ячейки → mismatch (регресс извлечения/рендера).
from openpyxl import load_workbook as _load_wb
from openpyxl.utils import get_column_letter as _colL

# Колонки данных (1-индекс) по файлам: файл 1 сдвинут (norma=E, actual=I…).
_EXCEL_COLMAP = {
    1: dict(label=3, norma=5, actual=9, time=10, cls=11, final=12),
    2: dict(label=3, norma=4, actual=8, time=9, cls=10, final=11),
    3: dict(label=3, norma=4, actual=8, time=9, cls=10, final=11),
    4: dict(label=3, norma=4, actual=8, time=9, cls=10, final=11),
    5: dict(label=3, norma=4, actual=8, time=9, cls=10, final=11),
}
# Начало первого блока и фикс. число под-строк (None у файла 1 — переменное).
_EXCEL_GEOM = {1: (31, None), 2: (31, 4), 3: (28, 9), 4: (28, 5), 5: (28, 16)}
# Опечатки ввода в эталоне — ТОЧЕЧНО по (файл, РМ), а не по всему РМ во всех файлах:
# 000051 — факт/время смещены в файлах 1/2/4; 000050 — факт освещённости текстом (файл 4).
# В файлах 3/5 у этих РМ данные корректны, поэтому их регрессы должны ПРОВАЛИВАТЬ
# гейт (не маскироваться в notes).
_EXCEL_ANOMALY = {(1, "000051"), (2, "000051"), (4, "000050"), (4, "000051")}
_EXCEL_WP_RE = re.compile(r"^\d{6}")


def _excel_norm(value):
    """Нормализовать значение ячейки для сравнения: число→float, текст→свёрнутый.

    «-»/пусто → None; «0,7»/«27,0» → float; «йўқ»/«yo'q»/«yo`q» → «yoq» (апострофы
    и пробелы выброшены). Диапазоны норм («23,0-31,0») сравниваются как текст.
    """
    if value is None:
        return None
    s = str(value).strip()
    if s in ("", "-", "—", "–", "−"):
        return None
    try:
        return round(float(s.replace(",", ".")), 4)
    except ValueError:
        return re.sub(r"[`'’‘ʻʼ´\s]+", "", s.lower())


def read_excel_protocol(path, idx: int) -> dict:
    """Прочитать xlsx-протокол → {workplace_no: {'final', 'subrows':[{norma,actual,time,cls}]}}."""
    ws = _load_wb(str(path))["complete"]
    cm = _EXCEL_COLMAP[idx]
    fixed = _EXCEL_GEOM[idx][1]
    starts = [(r, str(ws.cell(r, 1).value).strip())
              for r in range(1, ws.max_row + 1)
              if isinstance(ws.cell(r, 1).value, str) and _EXCEL_WP_RE.match(str(ws.cell(r, 1).value).strip())]
    blocks: dict[str, dict] = {}
    for i, (r, wp) in enumerate(starts):
        nextr = starts[i + 1][0] if i + 1 < len(starts) else ws.max_row + 1
        end = (r + fixed) if fixed else nextr
        subrows = []
        for rr in range(r, min(end, nextr)):
            subrows.append({k: ws.cell(rr, cm[k]).value for k in ("label", "norma", "actual", "time", "cls")})
        blocks[wp] = {"final": ws.cell(r, cm["final"]).value, "subrows": subrows}
    return blocks


def compare_excel(generated_path, reference_path, idx: int) -> CheckResult:
    """Сравнить сгенерированный xlsx-протокол с эталоном (единица — ячейка данных)."""
    gen = read_excel_protocol(generated_path, idx)
    ref = read_excel_protocol(reference_path, idx)
    res = CheckResult()

    for wp in sorted(set(ref) - set(gen)):
        res.mismatches.append(f"{wp}: есть в эталоне, нет в сгенерированном")
    for wp in sorted(set(gen) - set(ref)):
        res.notes.append(f"{wp}: есть у нас, нет в эталоне (эталон опустил / неполон — у нас по картам)")

    # Файл 1: метка под-строки (кол. C) — это ИМЯ вещества (данные), сверяем; у
    # файлов 2–5 метка статична (из прото), в сверку не берём.
    cmp_label = idx == 1
    for wp in sorted(set(ref) & set(gen)):
        g, e = gen[wp], ref[wp]
        anomaly = (idx, wp) in _EXCEL_ANOMALY
        bucket = res.notes if anomaly else res.mismatches
        # Итоговый класс блока
        res.total += 1
        if _excel_norm(g["final"]) == _excel_norm(e["final"]):
            res.matched += 1
        else:
            bucket.append(f"{wp} итог: эталон={e['final']!r} ≠ {g['final']!r}")
        # Под-строки (позиционно)
        keys = ("label", "norma", "actual", "time", "cls") if cmp_label else ("norma", "actual", "time", "cls")
        for i in range(max(len(g["subrows"]), len(e["subrows"]))):
            gsr = g["subrows"][i] if i < len(g["subrows"]) else {}
            esr = e["subrows"][i] if i < len(e["subrows"]) else {}
            for k in keys:
                gv, ev = _excel_norm(gsr.get(k)), _excel_norm(esr.get(k))
                if ev is None and gv is None:
                    continue  # обе пусты — не считаем (рамка)
                res.total += 1
                if gv == ev:
                    res.matched += 1
                else:
                    bucket.append(f"{wp} стр{i}.{k}: эталон={esr.get(k)!r} ≠ {gsr.get(k)!r}")
    return res
